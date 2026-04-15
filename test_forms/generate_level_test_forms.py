"""
自動回填分級測試表格產生器
產生 L1-L5 各級測試用 Excel/Word 檔案，用於驗證 form auto-fill pipeline。

使用方式：
    cd InduSpect/test_forms
    python generate_level_test_forms.py
"""

import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

# 修正 Windows 終端編碼
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# 共用樣式
THIN_BORDER = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)
TITLE_FONT = Font(name="標楷體", size=16, bold=True)
HEADER_FONT = Font(name="標楷體", size=11, bold=True)
NORMAL_FONT = Font(name="標楷體", size=10)
HEADER_FILL = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")


def apply_border_range(ws, min_row, max_row, min_col, max_col):
    """對指定範圍套用框線"""
    for row in range(min_row, max_row + 1):
        for col in range(min_col, max_col + 1):
            ws.cell(row=row, column=col).border = THIN_BORDER


# ================================================================
# Level 1：簡單 Key-Value 表
# ================================================================

def create_L1_excel():
    """
    L1_simple_kv.xlsx
    純 Key-Value 排列，標籤在左值在右，無合併格。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "基本資料"

    # 標題
    ws.merge_cells('A1:D1')
    ws['A1'] = "設備基本檢查表"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    # Key-Value 配對（標籤帶冒號，值格留空）
    fields = [
        ("設備名稱：", "", "設備編號：", ""),
        ("檢查日期：", "", "檢查人員：", ""),
        ("設備位置：", "", "廠區：", ""),
        ("溫度讀數：", "", "電壓值：", ""),
        ("轉速(RPM)：", "", "振動值：", ""),
        ("整體狀態：", "", "備註：", ""),
    ]

    for i, (k1, v1, k2, v2) in enumerate(fields, start=3):
        ws.cell(row=i, column=1, value=k1).font = HEADER_FONT
        ws.cell(row=i, column=2, value=v1).font = NORMAL_FONT
        ws.cell(row=i, column=3, value=k2).font = HEADER_FONT
        ws.cell(row=i, column=4, value=v2).font = NORMAL_FONT

    # 設定欄寬
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 20

    apply_border_range(ws, 3, 8, 1, 4)
    wb.save("L1_simple_kv.xlsx")
    print("✓ L1_simple_kv.xlsx")


def create_L1_word():
    """
    L1_simple_kv.docx
    Word 版 Key-Value，用表格呈現。
    """
    doc = Document()
    doc.add_heading("設備基本檢查表", level=1)

    table = doc.add_table(rows=6, cols=4, style='Table Grid')
    fields = [
        ("設備名稱：", "____", "設備編號：", "____"),
        ("檢查日期：", "____", "檢查人員：", "____"),
        ("設備位置：", "____", "廠區：", "____"),
        ("溫度讀數：", "____", "電壓值：", "____"),
        ("轉速(RPM)：", "____", "振動值：", "____"),
        ("整體狀態：", "____", "備註：", "____"),
    ]

    for i, (k1, v1, k2, v2) in enumerate(fields):
        cells = table.rows[i].cells
        cells[0].text = k1
        cells[1].text = v1
        cells[2].text = k2
        cells[3].text = v2
        # 標籤粗體
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
        for run in cells[2].paragraphs[0].runs:
            run.bold = True

    doc.save("L1_simple_kv.docx")
    print("✓ L1_simple_kv.docx")


# ================================================================
# Level 2：標準清單表
# ================================================================

def create_L2_standard_list():
    """
    L2_standard_list.xlsx
    固定表頭 + 逐行資料列。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "檢查清單"

    ws.merge_cells('A1:D1')
    ws['A1'] = "電氣設備定期檢查清單"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    # 表頭
    headers = ["檢查項目", "量測值", "標準範圍", "備註"]
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=3, column=j, value=h)
        c.font = HEADER_FONT
        c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal='center')

    # 資料列（值留空，等待回填）
    items = [
        ("絕緣電阻(MΩ)", "", "≥1.0", ""),
        ("接地電阻(Ω)", "", "≤10", ""),
        ("A相電流(A)", "", "10-15", ""),
        ("B相電流(A)", "", "10-15", ""),
        ("C相電流(A)", "", "10-15", ""),
        ("漏電流(mA)", "", "≤30", ""),
        ("電壓(V)", "", "380±10%", ""),
        ("溫度(℃)", "", "≤75", ""),
        ("噪音(dB)", "", "≤85", ""),
        ("振動(mm/s)", "", "≤4.5", ""),
    ]

    for i, (item, val, std, note) in enumerate(items, start=4):
        ws.cell(row=i, column=1, value=item).font = NORMAL_FONT
        ws.cell(row=i, column=2, value=val).font = NORMAL_FONT
        ws.cell(row=i, column=3, value=std).font = NORMAL_FONT
        ws.cell(row=i, column=4, value=note).font = NORMAL_FONT

    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 20

    apply_border_range(ws, 3, 13, 1, 4)
    wb.save("L2_standard_list.xlsx")
    print("✓ L2_standard_list.xlsx")


def create_L2_list_with_sections():
    """
    L2_list_with_sections.xlsx
    帶 section header 的清單（「一、電氣安全」等分段標題混在資料列中）。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "分段檢查表"

    ws.merge_cells('A1:D1')
    ws['A1'] = "設備安全定期檢查表"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    headers = ["檢查項目", "檢查結果", "標準值", "備註"]
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=3, column=j, value=h)
        c.font = HEADER_FONT
        c.fill = HEADER_FILL

    rows_data = [
        # section header（粗體，合併 4 欄）
        ("section", "一、電氣安全"),
        ("data", "絕緣電阻", "", "≥1.0 MΩ", ""),
        ("data", "接地電阻", "", "≤10 Ω", ""),
        ("data", "漏電斷路器", "", "正常動作", ""),
        # section header
        ("section", "二、機械安全"),
        ("data", "軸承溫度", "", "≤75℃", ""),
        ("data", "振動值", "", "≤4.5 mm/s", ""),
        ("data", "皮帶張力", "", "正常", ""),
        # section header
        ("section", "三、環境安全"),
        ("data", "噪音", "", "≤85 dB", ""),
        ("data", "粉塵量", "", "≤10 mg/m³", ""),
    ]

    row = 4
    for entry in rows_data:
        if entry[0] == "section":
            ws.merge_cells(f'A{row}:D{row}')
            c = ws.cell(row=row, column=1, value=entry[1])
            c.font = Font(name="標楷體", size=11, bold=True, color="2F5496")
            c.fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
        else:
            _, item, val, std, note = entry
            ws.cell(row=row, column=1, value=item).font = NORMAL_FONT
            ws.cell(row=row, column=2, value=val).font = NORMAL_FONT
            ws.cell(row=row, column=3, value=std).font = NORMAL_FONT
            ws.cell(row=row, column=4, value=note).font = NORMAL_FONT
        row += 1

    apply_border_range(ws, 3, row - 1, 1, 4)
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 20

    wb.save("L2_list_with_sections.xlsx")
    print("✓ L2_list_with_sections.xlsx")


# ================================================================
# Level 3：混合結構表
# ================================================================

def create_L3_mixed_excel():
    """
    L3_mixed_structure.xlsx
    上方 Key-Value + 中間清單 + 下方簽核。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "混合檢查表"

    # --- 標題 ---
    ws.merge_cells('A1:F1')
    ws['A1'] = "設備定期檢查紀錄表"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    # --- 上方：Key-Value 基本資料 ---
    kv_data = [
        ("設備名稱：", "", "", "設備編號：", "", ""),
        ("檢查日期：", "", "", "檢查人員：", "", ""),
        ("設備位置：", "", "", "所屬廠區：", "", ""),
    ]
    for i, row_data in enumerate(kv_data, start=3):
        for j, val in enumerate(row_data, 1):
            c = ws.cell(row=i, column=j, value=val)
            c.font = HEADER_FONT if j in (1, 4) else NORMAL_FONT

    apply_border_range(ws, 3, 5, 1, 6)

    # --- 中間：空一行後接清單 ---
    # 表頭
    list_headers = ["No.", "檢查項目", "量測值", "標準範圍", "判定", "備註"]
    for j, h in enumerate(list_headers, 1):
        c = ws.cell(row=7, column=j, value=h)
        c.font = HEADER_FONT
        c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal='center')

    list_items = [
        (1, "絕緣電阻(MΩ)", "", "≥1.0", "", ""),
        (2, "接地電阻(Ω)", "", "≤10", "", ""),
        (3, "A相電流(A)", "", "10-15", "", ""),
        (4, "B相電流(A)", "", "10-15", "", ""),
        (5, "表面溫度(℃)", "", "≤75", "", ""),
        (6, "振動(mm/s)", "", "≤4.5", "", ""),
    ]

    for i, (no, item, val, std, judge, note) in enumerate(list_items, start=8):
        ws.cell(row=i, column=1, value=no).font = NORMAL_FONT
        ws.cell(row=i, column=2, value=item).font = NORMAL_FONT
        ws.cell(row=i, column=3, value=val).font = NORMAL_FONT
        ws.cell(row=i, column=4, value=std).font = NORMAL_FONT
        ws.cell(row=i, column=5, value=judge).font = NORMAL_FONT
        ws.cell(row=i, column=6, value=note).font = NORMAL_FONT

    apply_border_range(ws, 7, 13, 1, 6)

    # --- 下方：簽核欄 ---
    ws.cell(row=15, column=1, value="檢查人員簽章：").font = HEADER_FONT
    ws.cell(row=15, column=2, value="").font = NORMAL_FONT
    ws.cell(row=15, column=4, value="主管簽章：").font = HEADER_FONT
    ws.cell(row=15, column=5, value="").font = NORMAL_FONT
    ws.cell(row=16, column=1, value="簽核日期：").font = HEADER_FONT
    ws.cell(row=16, column=2, value="").font = NORMAL_FONT
    apply_border_range(ws, 15, 16, 1, 6)

    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws.column_dimensions[col].width = 15

    wb.save("L3_mixed_structure.xlsx")
    print("✓ L3_mixed_structure.xlsx")


def create_L3_mixed_word():
    """
    L3_mixed_structure.docx
    上方段落 Key-Value + 中間表格清單 + 下方簽核段落。
    """
    doc = Document()
    doc.add_heading("設備定期檢查紀錄表", level=1)

    # --- 上方：段落式 Key-Value ---
    for label, placeholder in [
        ("設備名稱：", "________________"),
        ("設備編號：", "________________"),
        ("檢查日期：", "________________"),
        ("檢查人員：", "________________"),
        ("設備位置：", "________________"),
    ]:
        p = doc.add_paragraph()
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_val = p.add_run(placeholder)
        run_val.font.size = Pt(12)

    doc.add_paragraph()  # 空行

    # --- 中間：表格清單 ---
    doc.add_heading("檢查項目", level=2)
    table = doc.add_table(rows=7, cols=4, style='Table Grid')
    headers = ["檢查項目", "量測值", "標準範圍", "備註"]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    items = [
        ("絕緣電阻(MΩ)", "____", "≥1.0", ""),
        ("接地電阻(Ω)", "____", "≤10", ""),
        ("A相電流(A)", "____", "10-15", ""),
        ("B相電流(A)", "____", "10-15", ""),
        ("表面溫度(℃)", "____", "≤75", ""),
        ("振動(mm/s)", "____", "≤4.5", ""),
    ]
    for i, (item, val, std, note) in enumerate(items, start=1):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = val
        table.rows[i].cells[2].text = std
        table.rows[i].cells[3].text = note

    doc.add_paragraph()

    # --- 下方：簽核 ---
    doc.add_heading("簽核", level=2)
    p = doc.add_paragraph()
    p.add_run("檢查人員簽章：").bold = True
    p.add_run("________________")
    p = doc.add_paragraph()
    p.add_run("主管簽章：").bold = True
    p.add_run("________________")
    p = doc.add_paragraph()
    p.add_run("簽核日期：").bold = True
    p.add_run("________________")

    doc.save("L3_mixed_structure.docx")
    print("✓ L3_mixed_structure.docx")


# ================================================================
# Level 4：合併儲存格 + 雙欄 Checkbox
# ================================================================

def create_L4_merged_cells():
    """
    L4_merged_cells.xlsx
    大量合併儲存格的表格。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "合併格測試"

    # 標題：跨 6 欄合併
    ws.merge_cells('A1:F1')
    ws['A1'] = "合併儲存格測試表"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    # Key-Value 但標籤跨 2 欄合併
    # Row 3: 「設備名稱」合併 A3:B3，值在 C3:D3 合併
    ws.merge_cells('A3:B3')
    ws['A3'] = "設備名稱："
    ws['A3'].font = HEADER_FONT
    ws.merge_cells('C3:D3')
    ws['C3'] = ""  # 待填

    ws.merge_cells('E3:E3')  # 不合併，單格
    ws['E3'] = "編號："
    ws['E3'].font = HEADER_FONT
    ws['F3'] = ""  # 待填

    # Row 4: 標籤跨 2 欄
    ws.merge_cells('A4:B4')
    ws['A4'] = "檢查日期："
    ws['A4'].font = HEADER_FONT
    ws.merge_cells('C4:D4')
    ws['C4'] = ""  # 待填

    ws.merge_cells('E4:F4')
    ws['E4'] = ""  # 故意空的合併格

    # Row 5: 標籤跨 3 欄合併
    ws.merge_cells('A5:C5')
    ws['A5'] = "設備安裝位置與廠區說明："
    ws['A5'].font = HEADER_FONT
    ws.merge_cells('D5:F5')
    ws['D5'] = ""  # 值格跨 3 欄

    # Row 7: 跨行合併的標籤
    ws.merge_cells('A7:A8')
    ws['A7'] = "電氣\n參數"
    ws['A7'].font = HEADER_FONT
    ws['A7'].alignment = Alignment(wrap_text=True, vertical='center')

    ws['B7'] = "電壓(V)："
    ws['B7'].font = NORMAL_FONT
    ws['C7'] = ""  # 待填
    ws['B8'] = "電流(A)："
    ws['B8'].font = NORMAL_FONT
    ws['C8'] = ""  # 待填

    ws.merge_cells('D7:D8')
    ws['D7'] = "機械\n參數"
    ws['D7'].font = HEADER_FONT
    ws['D7'].alignment = Alignment(wrap_text=True, vertical='center')

    ws['E7'] = "溫度(℃)："
    ws['E7'].font = NORMAL_FONT
    ws['F7'] = ""  # 待填
    ws['E8'] = "振動(mm/s)："
    ws['E8'].font = NORMAL_FONT
    ws['F8'] = ""  # 待填

    apply_border_range(ws, 3, 8, 1, 6)
    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws.column_dimensions[col].width = 15

    wb.save("L4_merged_cells.xlsx")
    print("✓ L4_merged_cells.xlsx")


def create_L4_dual_checkbox():
    """
    L4_dual_checkbox.xlsx
    合格/不合格雙欄 checkbox 表格。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "雙欄勾選"

    ws.merge_cells('A1:E1')
    ws['A1'] = "設備定期檢查表（勾選式）"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    # 表頭
    headers = ["No.", "檢查項目", "合格", "不合格", "備註"]
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=3, column=j, value=h)
        c.font = HEADER_FONT
        c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal='center')

    # 檢查項目（合格/不合格欄留空等待填入 ✓）
    items = [
        (1, "絕緣測試", "☐", "☐", ""),
        (2, "接地測試", "☐", "☐", ""),
        (3, "漏電斷路器動作", "☐", "☐", ""),
        (4, "線路絕緣狀態", "☐", "☐", ""),
        (5, "電氣接點狀態", "☐", "☐", ""),
        (6, "防爆設備完整性", "☐", "☐", ""),
        (7, "過載保護裝置", "☐", "☐", ""),
        (8, "緊急停止裝置", "☐", "☐", ""),
    ]

    for i, (no, item, ok, ng, note) in enumerate(items, start=4):
        ws.cell(row=i, column=1, value=no).font = NORMAL_FONT
        ws.cell(row=i, column=1).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=2, value=item).font = NORMAL_FONT
        ws.cell(row=i, column=3, value=ok).font = NORMAL_FONT
        ws.cell(row=i, column=3).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=4, value=ng).font = NORMAL_FONT
        ws.cell(row=i, column=4).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=5, value=note).font = NORMAL_FONT

    apply_border_range(ws, 3, 11, 1, 5)
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 20

    wb.save("L4_dual_checkbox.xlsx")
    print("✓ L4_dual_checkbox.xlsx")


def create_L4_merged_checkbox_combo():
    """
    L4_merged_checkbox_combo.xlsx
    合併儲存格 + 雙欄 checkbox 的複合場景。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "合併勾選複合"

    ws.merge_cells('A1:F1')
    ws['A1'] = "設備安全檢查表（複合格式）"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    # Key-Value 區（有合併格）
    ws.merge_cells('A3:B3')
    ws['A3'] = "設備名稱："
    ws['A3'].font = HEADER_FONT
    ws.merge_cells('C3:D3')
    ws['C3'] = ""

    ws.merge_cells('E3:E3')
    ws['E3'] = "日期："
    ws['E3'].font = HEADER_FONT
    ws['F3'] = ""

    # 分段標題（合併全行）
    ws.merge_cells('A5:F5')
    ws['A5'] = "一、安全檢查項目"
    ws['A5'].font = Font(name="標楷體", size=11, bold=True, color="2F5496")
    ws['A5'].fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")

    # checkbox 表頭
    headers = ["類別", "檢查項目", "合格", "不合格", "量測值", "備註"]
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=6, column=j, value=h)
        c.font = HEADER_FONT
        c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal='center')

    # 類別合併 + checkbox
    # 「電氣安全」跨 3 行合併
    ws.merge_cells('A7:A9')
    ws['A7'] = "電氣安全"
    ws['A7'].font = HEADER_FONT
    ws['A7'].alignment = Alignment(vertical='center', wrap_text=True)

    elec_items = [
        ("絕緣電阻", "☐", "☐", "", ""),
        ("接地電阻", "☐", "☐", "", ""),
        ("漏電保護", "☐", "☐", "", ""),
    ]
    for i, (item, ok, ng, val, note) in enumerate(elec_items, start=7):
        ws.cell(row=i, column=2, value=item).font = NORMAL_FONT
        ws.cell(row=i, column=3, value=ok).font = NORMAL_FONT
        ws.cell(row=i, column=3).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=4, value=ng).font = NORMAL_FONT
        ws.cell(row=i, column=4).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=5, value=val).font = NORMAL_FONT
        ws.cell(row=i, column=6, value=note).font = NORMAL_FONT

    # 「機械安全」跨 2 行合併
    ws.merge_cells('A10:A11')
    ws['A10'] = "機械安全"
    ws['A10'].font = HEADER_FONT
    ws['A10'].alignment = Alignment(vertical='center', wrap_text=True)

    mech_items = [
        ("軸承狀態", "☐", "☐", "", ""),
        ("皮帶張力", "☐", "☐", "", ""),
    ]
    for i, (item, ok, ng, val, note) in enumerate(mech_items, start=10):
        ws.cell(row=i, column=2, value=item).font = NORMAL_FONT
        ws.cell(row=i, column=3, value=ok).font = NORMAL_FONT
        ws.cell(row=i, column=3).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=4, value=ng).font = NORMAL_FONT
        ws.cell(row=i, column=4).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=5, value=val).font = NORMAL_FONT
        ws.cell(row=i, column=6, value=note).font = NORMAL_FONT

    apply_border_range(ws, 3, 11, 1, 6)
    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws.column_dimensions[col].width = 14

    wb.save("L4_merged_checkbox_combo.xlsx")
    print("✓ L4_merged_checkbox_combo.xlsx")


# ================================================================
# Level 5：自由格式 / 非標準表
# ================================================================

def create_L5_nonstandard_labels():
    """
    L5_nonstandard_labels.xlsx
    欄位名用英文、簡寫、非標準用語。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "NonStd"

    ws.merge_cells('A1:D1')
    ws['A1'] = "Equipment Periodic Inspection Form"
    ws['A1'].font = Font(size=14, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')

    # 非標準命名的 Key-Value
    fields = [
        ("Equip. Name:", "", "S/N:", ""),
        ("Insp. Date:", "", "Inspector:", ""),
        ("Location:", "", "Dept:", ""),
    ]
    for i, (k1, v1, k2, v2) in enumerate(fields, start=3):
        ws.cell(row=i, column=1, value=k1).font = Font(size=10, bold=True)
        ws.cell(row=i, column=2).font = Font(size=10)
        ws.cell(row=i, column=3, value=k2).font = Font(size=10, bold=True)
        ws.cell(row=i, column=4).font = Font(size=10)

    # 非標準清單
    headers = ["Item", "Reading", "Spec", "P/F", "Rmks"]
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=6, column=j, value=h)
        c.font = Font(size=10, bold=True)
        c.fill = HEADER_FILL

    items = [
        ("Ins. Res.", "", "≥1.0MΩ", "", ""),
        ("Gnd. Res.", "", "≤10Ω", "", ""),
        ("Temp.", "", "≤75℃", "", ""),
        ("Vib.", "", "≤4.5mm/s", "", ""),
        ("RPM", "", "1450±50", "", ""),
    ]
    for i, row_data in enumerate(items, start=7):
        for j, val in enumerate(row_data, 1):
            ws.cell(row=i, column=j, value=val).font = Font(size=10)

    apply_border_range(ws, 3, 11, 1, 5)
    wb.save("L5_nonstandard_labels.xlsx")
    print("✓ L5_nonstandard_labels.xlsx")


def create_L5_oversized():
    """
    L5_oversized.xlsx
    超過 200 行的超長表格，測試截斷處理。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "超長表格"

    ws['A1'] = "超大設備清單檢查表"
    ws['A1'].font = TITLE_FONT

    headers = ["No.", "設備名稱", "檢查結果", "備註"]
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=2, column=j, value=h)
        c.font = HEADER_FONT

    # 產生 300 行資料
    for i in range(3, 303):
        ws.cell(row=i, column=1, value=i - 2)
        ws.cell(row=i, column=2, value=f"設備_{i-2:03d}")
        ws.cell(row=i, column=3, value="")  # 待填
        ws.cell(row=i, column=4, value="")

    wb.save("L5_oversized.xlsx")
    print("✓ L5_oversized.xlsx")


def create_L5_multi_sheet():
    """
    L5_multi_sheet.xlsx
    多 sheet（含一個空白 sheet）的 Excel。
    """
    wb = openpyxl.Workbook()

    # Sheet1: 基本資料
    ws1 = wb.active
    ws1.title = "基本資料"
    ws1['A1'] = "設備名稱："
    ws1['A1'].font = HEADER_FONT
    ws1['B1'] = ""
    ws1['A2'] = "檢查日期："
    ws1['A2'].font = HEADER_FONT
    ws1['B2'] = ""

    # Sheet2: 空白（故意的）
    wb.create_sheet("空白頁")

    # Sheet3: 檢查項目
    ws3 = wb.create_sheet("檢查項目")
    headers = ["項目", "結果", "備註"]
    for j, h in enumerate(headers, 1):
        ws3.cell(row=1, column=j, value=h).font = HEADER_FONT

    for i in range(2, 8):
        ws3.cell(row=i, column=1, value=f"檢查項_{i-1}")
        ws3.cell(row=i, column=2, value="")
        ws3.cell(row=i, column=3, value="")

    # Sheet4: 同名欄位
    ws4 = wb.create_sheet("複檢紀錄")
    ws4['A1'] = "設備名稱："  # 與 Sheet1 同名
    ws4['A1'].font = HEADER_FONT
    ws4['B1'] = ""
    ws4['A2'] = "複檢日期："
    ws4['A2'].font = HEADER_FONT
    ws4['B2'] = ""
    ws4['A3'] = "複檢結果："
    ws4['A3'].font = HEADER_FONT
    ws4['B3'] = ""

    wb.save("L5_multi_sheet.xlsx")
    print("✓ L5_multi_sheet.xlsx")


def create_L5_nested_table_word():
    """
    L5_nested_table.docx
    Word 嵌套表格（表中表）。
    """
    doc = Document()
    doc.add_heading("設備檢查表（嵌套格式）", level=1)

    # 外層表格
    outer = doc.add_table(rows=4, cols=2, style='Table Grid')

    outer.rows[0].cells[0].text = "設備名稱："
    outer.rows[0].cells[1].text = "____"
    outer.rows[1].cells[0].text = "檢查日期："
    outer.rows[1].cells[1].text = "____"

    # 第 3 行：左格放文字，右格放「內嵌表格」
    outer.rows[2].cells[0].text = "詳細檢查紀錄"

    # 在右格內插入嵌套表格
    inner_cell = outer.rows[2].cells[1]
    inner_cell.text = ""  # 清空
    inner_table = inner_cell.add_table(rows=4, cols=3)
    inner_table.style = 'Table Grid'

    inner_headers = ["項目", "結果", "備註"]
    for j, h in enumerate(inner_headers):
        inner_table.rows[0].cells[j].text = h

    inner_items = [
        ("電壓", "____", ""),
        ("電流", "____", ""),
        ("溫度", "____", ""),
    ]
    for i, (item, val, note) in enumerate(inner_items, start=1):
        inner_table.rows[i].cells[0].text = item
        inner_table.rows[i].cells[1].text = val
        inner_table.rows[i].cells[2].text = note

    # 第 4 行：簽核
    outer.rows[3].cells[0].text = "簽核："
    outer.rows[3].cells[1].text = "____"

    doc.save("L5_nested_table.docx")
    print("✓ L5_nested_table.docx")


def create_L5_irregular_merge():
    """
    L5_irregular_merge.xlsx
    不規則合併儲存格（非矩形感覺的複雜佈局）。
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "不規則合併"

    ws.merge_cells('A1:H1')
    ws['A1'] = "複雜格式設備巡檢表"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = Alignment(horizontal='center')

    # 第一段：不同大小的合併格交錯
    ws.merge_cells('A3:C3')
    ws['A3'] = "設備基本資訊"
    ws['A3'].font = Font(size=12, bold=True)

    ws.merge_cells('A4:A5')  # 跨行
    ws['A4'] = "名稱/\n編號"
    ws['A4'].alignment = Alignment(wrap_text=True, vertical='center')
    ws.merge_cells('B4:C4')
    ws['B4'] = ""  # 設備名稱值
    ws.merge_cells('B5:C5')
    ws['B5'] = ""  # 設備編號值

    ws.merge_cells('D3:H3')
    ws['D3'] = "檢查資訊"
    ws['D3'].font = Font(size=12, bold=True)

    ws.merge_cells('D4:E4')
    ws['D4'] = "日期："
    ws.merge_cells('F4:H4')
    ws['F4'] = ""
    ws.merge_cells('D5:E5')
    ws['D5'] = "人員："
    ws.merge_cells('F5:H5')
    ws['F5'] = ""

    # 第二段：更複雜的合併
    ws.merge_cells('A7:H7')
    ws['A7'] = "檢測數據"
    ws['A7'].font = Font(size=12, bold=True)
    ws['A7'].fill = HEADER_FILL

    # 多層表頭
    ws.merge_cells('A8:A9')
    ws['A8'] = "No."
    ws.merge_cells('B8:B9')
    ws['B8'] = "項目"
    ws.merge_cells('C8:D8')
    ws['C8'] = "量測值"
    ws['C9'] = "數值"
    ws['D9'] = "單位"
    ws.merge_cells('E8:F8')
    ws['E8'] = "判定"
    ws['E9'] = "合格"
    ws['F9'] = "不合格"
    ws.merge_cells('G8:H8')
    ws['G8'] = "標準"
    ws['G9'] = "下限"
    ws['H9'] = "上限"

    # 資料列
    data = [
        (1, "絕緣電阻", "", "MΩ", "☐", "☐", "1.0", "∞"),
        (2, "接地電阻", "", "Ω", "☐", "☐", "0", "10"),
        (3, "溫度", "", "℃", "☐", "☐", "0", "75"),
    ]
    for i, (no, item, val, unit, ok, ng, lo, hi) in enumerate(data, start=10):
        ws.cell(row=i, column=1, value=no)
        ws.cell(row=i, column=2, value=item)
        ws.cell(row=i, column=3, value=val)
        ws.cell(row=i, column=4, value=unit)
        ws.cell(row=i, column=5, value=ok).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=6, value=ng).alignment = Alignment(horizontal='center')
        ws.cell(row=i, column=7, value=lo)
        ws.cell(row=i, column=8, value=hi)

    apply_border_range(ws, 3, 12, 1, 8)
    for col in range(1, 9):
        ws.column_dimensions[get_column_letter(col)].width = 12

    wb.save("L5_irregular_merge.xlsx")
    print("✓ L5_irregular_merge.xlsx")


# ================================================================
# 主程式
# ================================================================

if __name__ == "__main__":
    # 確保在 test_forms 目錄執行
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)

    print("=== 產生 L1-L5 分級測試表格 ===\n")

    # Level 1
    create_L1_excel()
    create_L1_word()

    # Level 2
    create_L2_standard_list()
    create_L2_list_with_sections()

    # Level 3
    create_L3_mixed_excel()
    create_L3_mixed_word()

    # Level 4
    create_L4_merged_cells()
    create_L4_dual_checkbox()
    create_L4_merged_checkbox_combo()

    # Level 5
    create_L5_nonstandard_labels()
    create_L5_oversized()
    create_L5_multi_sheet()
    create_L5_nested_table_word()
    create_L5_irregular_merge()

    print(f"\n=== 完成！共產生 14 個測試檔案 ===")
