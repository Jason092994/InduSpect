"""
建立測試用定檢表單 - 模擬真實廠商/學校的設備定期檢查表格
"""

import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_electrical_inspection_excel():
    """
    建立：電氣設備定期檢查表 (Excel)
    模擬北科大環安中心的真實檢查表格式
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "電氣設備定期檢查表"

    # 樣式
    title_font = Font(name="標楷體", size=16, bold=True)
    header_font = Font(name="標楷體", size=11, bold=True)
    normal_font = Font(name="標楷體", size=10)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # 設定欄寬
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 20

    # 標題
    ws.merge_cells('A1:G1')
    ws['A1'] = '國立臺北科技大學 電氣設備定期檢查表'
    ws['A1'].font = title_font
    ws['A1'].alignment = center_align

    # 表單編號列
    ws.merge_cells('A2:C2')
    ws['A2'] = '表單編號：EI-2024-'
    ws['A2'].font = normal_font
    ws.merge_cells('D2:G2')
    ws['D2'] = '檢查週期：□每月 □每季 □每半年 □每年'
    ws['D2'].font = normal_font

    # 基本資訊區塊
    row = 3
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '一、基本資訊'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill

    info_fields = [
        ('設備名稱', '', '設備編號', ''),
        ('設備位置', '', '額定電壓/電流', ''),
        ('製造廠商', '', '安裝日期', ''),
        ('檢查日期', '', '檢查人員', ''),
        ('陪同人員', '', '天氣狀況', ''),
    ]

    for i, (label1, val1, label2, val2) in enumerate(info_fields):
        r = row + 1 + i
        ws[f'A{r}'] = ''
        ws[f'B{r}'] = label1
        ws[f'B{r}'].font = header_font
        ws[f'C{r}'] = val1
        ws[f'C{r}'].font = normal_font
        ws[f'D{r}'] = ''
        ws[f'E{r}'] = label2
        ws[f'E{r}'].font = header_font
        ws[f'F{r}'] = val2
        ws[f'F{r}'].font = normal_font
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws[f'{col}{r}'].border = thin_border

    # 檢查項目區塊
    row = 9
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '二、檢查項目'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill

    # 檢查項目表頭
    row = 10
    headers = ['項次', '檢查項目', '檢查標準', '正常', '異常', '不適用', '備註/異常說明']
    for i, h in enumerate(headers):
        col = chr(ord('A') + i)
        ws[f'{col}{row}'] = h
        ws[f'{col}{row}'].font = header_font
        ws[f'{col}{row}'].alignment = center_align
        ws[f'{col}{row}'].border = thin_border
        ws[f'{col}{row}'].fill = header_fill

    # 檢查項目內容
    items = [
        ('1', '配電盤外觀', '無鏽蝕、變形、破損'),
        ('2', '配電盤門鎖', '門鎖功能正常，可正常開關'),
        ('3', '接地線連接', '接地線連接牢固，無鬆脫'),
        ('4', '電纜絕緣', '絕緣層完整，無老化龜裂'),
        ('5', '斷路器動作', '手動操作正常，跳脫功能正常'),
        ('6', '指示燈狀態', '各指示燈正常顯示'),
        ('7', '電壓量測', '電壓值在額定範圍±10%內'),
        ('8', '電流量測', '電流值未超過額定值'),
        ('9', '溫度檢測', '各接點溫度正常，無異常發熱'),
        ('10', '漏電檢測', '漏電電流<30mA'),
        ('11', '接線端子', '端子鎖固正常，無鬆動變色'),
        ('12', '通風散熱', '散熱風扇運轉正常，通風口無阻塞'),
        ('13', '標示標籤', '各回路標示清楚，警告標誌完整'),
        ('14', '防護設備', '絕緣手套、絕緣毯等防護設備齊全'),
    ]

    for i, (num, item, standard) in enumerate(items):
        r = row + 1 + i
        ws[f'A{r}'] = num
        ws[f'A{r}'].alignment = center_align
        ws[f'B{r}'] = item
        ws[f'C{r}'] = standard
        ws[f'C{r}'].alignment = left_align
        ws[f'D{r}'] = '□'
        ws[f'D{r}'].alignment = center_align
        ws[f'E{r}'] = '□'
        ws[f'E{r}'].alignment = center_align
        ws[f'F{r}'] = '□'
        ws[f'F{r}'].alignment = center_align
        ws[f'G{r}'] = ''
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws[f'{col}{r}'].border = thin_border
            ws[f'{col}{r}'].font = normal_font

    # 量測數據區塊
    row = 25
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '三、量測數據'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill

    measure_headers = ['項次', '量測項目', '量測位置', '量測值', '單位', '標準值', '判定']
    row = 26
    for i, h in enumerate(measure_headers):
        col = chr(ord('A') + i)
        ws[f'{col}{row}'] = h
        ws[f'{col}{row}'].font = header_font
        ws[f'{col}{row}'].alignment = center_align
        ws[f'{col}{row}'].border = thin_border
        ws[f'{col}{row}'].fill = header_fill

    measures = [
        ('1', 'R相電壓', '主配電盤', '', 'V', '220±22V', ''),
        ('2', 'S相電壓', '主配電盤', '', 'V', '220±22V', ''),
        ('3', 'T相電壓', '主配電盤', '', 'V', '220±22V', ''),
        ('4', 'R相電流', '主配電盤', '', 'A', '<額定值', ''),
        ('5', 'S相電流', '主配電盤', '', 'A', '<額定值', ''),
        ('6', 'T相電流', '主配電盤', '', 'A', '<額定值', ''),
        ('7', '接地電阻', '接地極', '', 'Ω', '<10Ω', ''),
        ('8', '絕緣電阻', '主迴路', '', 'MΩ', '>1MΩ', ''),
    ]

    for i, (num, item, loc, val, unit, std, judge) in enumerate(measures):
        r = row + 1 + i
        ws[f'A{r}'] = num
        ws[f'A{r}'].alignment = center_align
        ws[f'B{r}'] = item
        ws[f'C{r}'] = loc
        ws[f'D{r}'] = val
        ws[f'E{r}'] = unit
        ws[f'E{r}'].alignment = center_align
        ws[f'F{r}'] = std
        ws[f'G{r}'] = judge
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws[f'{col}{r}'].border = thin_border
            ws[f'{col}{r}'].font = normal_font

    # 結論區塊
    row = 35
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '四、綜合判定與建議'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill

    row = 36
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '綜合判定：□合格  □有條件合格  □不合格'
    ws[f'A{row}'].font = normal_font
    ws[f'A{row}'].border = thin_border

    row = 37
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '改善建議：'
    ws[f'A{row}'].font = normal_font
    ws[f'A{row}'].border = thin_border

    row = 38
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = ''
    ws[f'A{row}'].border = thin_border
    ws.row_dimensions[row].height = 60

    row = 39
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '預計改善完成日期：     年     月     日'
    ws[f'A{row}'].font = normal_font
    ws[f'A{row}'].border = thin_border

    # 簽核區
    row = 41
    ws.merge_cells(f'A{row}:G{row}')
    ws[f'A{row}'] = '五、簽核'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill

    row = 42
    sign_headers = ['', '檢查人員', '', '單位主管', '', '環安中心', '']
    for i, h in enumerate(sign_headers):
        col = chr(ord('A') + i)
        ws[f'{col}{row}'] = h
        ws[f'{col}{row}'].font = header_font
        ws[f'{col}{row}'].alignment = center_align
        ws[f'{col}{row}'].border = thin_border

    row = 43
    ws.merge_cells(f'A{row}:A{row}')
    ws[f'A{row}'] = '簽名'
    ws[f'A{row}'].font = normal_font
    ws.merge_cells(f'B{row}:C{row}')
    ws[f'B{row}'] = ''
    ws.merge_cells(f'D{row}:E{row}')
    ws[f'D{row}'] = ''
    ws.merge_cells(f'F{row}:G{row}')
    ws[f'F{row}'] = ''
    for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws[f'{col}{row}'].border = thin_border
    ws.row_dimensions[row].height = 40

    row = 44
    ws[f'A{row}'] = '日期'
    ws[f'A{row}'].font = normal_font
    ws.merge_cells(f'B{row}:C{row}')
    ws[f'B{row}'] = ''
    ws.merge_cells(f'D{row}:E{row}')
    ws[f'D{row}'] = ''
    ws.merge_cells(f'F{row}:G{row}')
    ws[f'F{row}'] = ''
    for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws[f'{col}{row}'].border = thin_border

    output = '/home/user/InduSpect/test_forms/電氣設備定期檢查表.xlsx'
    wb.save(output)
    print(f'✅ Created: {output}')
    return output


def create_fire_safety_inspection_docx():
    """
    建立：消防安全設備定期檢查表 (Word)
    模擬真實的消防設備巡檢表格式
    """
    doc = Document()

    # 設定預設字型
    style = doc.styles['Normal']
    font = style.font
    font.name = '標楷體'
    font.size = Pt(11)

    # 標題
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('國立臺北科技大學\n消防安全設備定期檢查紀錄表')
    run.bold = True
    run.font.size = Pt(16)

    # 表單資訊
    info_para = doc.add_paragraph()
    info_para.add_run('表單編號：FS-2024-          ').bold = False
    info_para.add_run('修訂版次：第3版').bold = False

    # 基本資訊表格
    doc.add_paragraph('一、基本資訊', style='Heading 2')

    table1 = doc.add_table(rows=4, cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ('建築物名稱', '', '樓層/區域', ''),
        ('檢查日期', '   年   月   日', '檢查時間', '   :   ~   :   '),
        ('檢查人員', '', '證照號碼', ''),
        ('陪同人員', '', '聯絡電話', ''),
    ]

    for i, (l1, v1, l2, v2) in enumerate(info_data):
        row = table1.rows[i]
        row.cells[0].text = l1
        row.cells[1].text = v1
        row.cells[2].text = l2
        row.cells[3].text = v2
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 滅火器檢查
    doc.add_paragraph('二、滅火器檢查', style='Heading 2')

    table2 = doc.add_table(rows=9, cols=6)
    table2.style = 'Table Grid'

    headers2 = ['項次', '檢查項目', '檢查要點', '合格', '不合格', '備註']
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
        for paragraph in table2.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    fire_items = [
        ('1', '設置位置', '設置於明顯且易於取用之處'),
        ('2', '標示', '有明顯標示，標示清晰可辨'),
        ('3', '外觀', '本體無變形、損傷、腐蝕'),
        ('4', '壓力錶', '指針在綠色範圍內'),
        ('5', '安全插銷', '安全插銷完整、封條未拆'),
        ('6', '噴嘴/皮管', '噴嘴暢通、皮管無龜裂'),
        ('7', '有效期限', '在有效期限內'),
        ('8', '放置高度', '頂端離地面≤1.5公尺'),
    ]

    for i, (num, item, point) in enumerate(fire_items):
        row = table2.rows[i + 1]
        row.cells[0].text = num
        row.cells[1].text = item
        row.cells[2].text = point
        row.cells[3].text = '□'
        row.cells[4].text = '□'
        row.cells[5].text = ''

    doc.add_paragraph()

    # 室內消防栓檢查
    doc.add_paragraph('三、室內消防栓檢查', style='Heading 2')

    table3 = doc.add_table(rows=8, cols=6)
    table3.style = 'Table Grid'

    for i, h in enumerate(headers2):
        table3.rows[0].cells[i].text = h
        for paragraph in table3.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    hydrant_items = [
        ('1', '箱體外觀', '箱門開閉正常、玻璃完整'),
        ('2', '消防栓閥', '開關操作靈活、無漏水'),
        ('3', '水帶', '水帶完整、無破損發霉'),
        ('4', '瞄子', '瞄子完整、無阻塞'),
        ('5', '標示燈', '標示燈正常亮起'),
        ('6', '加壓泵浦', '啟動正常、壓力正常'),
        ('7', '送水口', '送水口蓋完整、無阻塞'),
    ]

    for i, (num, item, point) in enumerate(hydrant_items):
        row = table3.rows[i + 1]
        row.cells[0].text = num
        row.cells[1].text = item
        row.cells[2].text = point
        row.cells[3].text = '□'
        row.cells[4].text = '□'
        row.cells[5].text = ''

    doc.add_paragraph()

    # 火災警報設備
    doc.add_paragraph('四、火災警報設備檢查', style='Heading 2')

    table4 = doc.add_table(rows=7, cols=6)
    table4.style = 'Table Grid'

    for i, h in enumerate(headers2):
        table4.rows[0].cells[i].text = h
        for paragraph in table4.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    alarm_items = [
        ('1', '偵煙探測器', '外觀正常、指示燈正常閃爍'),
        ('2', '定溫探測器', '外觀正常、無遮蔽物'),
        ('3', '手動報警機', '按鈕正常、保護蓋完整'),
        ('4', '受信總機', '電源正常、無故障警示'),
        ('5', '警鈴/蜂鳴器', '音量足夠、動作正常'),
        ('6', '緊急廣播', '播放正常、音量適當'),
    ]

    for i, (num, item, point) in enumerate(alarm_items):
        row = table4.rows[i + 1]
        row.cells[0].text = num
        row.cells[1].text = item
        row.cells[2].text = point
        row.cells[3].text = '□'
        row.cells[4].text = '□'
        row.cells[5].text = ''

    doc.add_paragraph()

    # 綜合評估
    doc.add_paragraph('五、綜合評估', style='Heading 2')

    eval_table = doc.add_table(rows=5, cols=2)
    eval_table.style = 'Table Grid'

    eval_data = [
        ('整體評估結果', '□全部合格  □部分不合格  □嚴重不合格'),
        ('不合格項目數', '       項'),
        ('改善建議', '\n\n\n'),
        ('預計改善日期', '   年   月   日'),
        ('複查日期', '   年   月   日'),
    ]

    for i, (label, value) in enumerate(eval_data):
        eval_table.rows[i].cells[0].text = label
        eval_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 簽核
    doc.add_paragraph('六、簽核', style='Heading 2')

    sign_table = doc.add_table(rows=3, cols=4)
    sign_table.style = 'Table Grid'

    sign_headers = ['', '檢查人員', '單位主管', '消防管理人']
    for i, h in enumerate(sign_headers):
        sign_table.rows[0].cells[i].text = h

    sign_table.rows[1].cells[0].text = '簽名'
    sign_table.rows[2].cells[0].text = '日期'

    # 注意事項
    doc.add_paragraph()
    note = doc.add_paragraph()
    run = note.add_run('注意事項：')
    run.bold = True
    doc.add_paragraph('1. 本表依據「各類場所消防安全設備設置標準」及「消防安全設備檢修及申報辦法」辦理。')
    doc.add_paragraph('2. 檢查頻率：每月至少一次外觀檢查，每半年一次綜合檢查。')
    doc.add_paragraph('3. 發現不合格項目應立即通報環安中心並限期改善。')
    doc.add_paragraph('4. 本表填寫完畢後請送環安中心備查存檔。')

    output = '/home/user/InduSpect/test_forms/消防安全設備定期檢查表.docx'
    doc.save(output)
    print(f'✅ Created: {output}')
    return output


if __name__ == '__main__':
    create_electrical_inspection_excel()
    create_fire_safety_inspection_docx()
    print('\n🎉 All test forms created successfully!')
