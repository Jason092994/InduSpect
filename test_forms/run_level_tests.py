"""
表單自動回填分級測試 — 完整測試 + HTML 報告產生

對 L1-L5 共 14 個測試表格執行：
1. 結構分析（欄位偵測）
2. 模擬回填（使用 mock 資料）
3. 讀回填入後的值
4. 產出 before/after 對照的完整 HTML 報告

使用方式：
    cd InduSpect/test_forms
    python run_level_tests.py
"""

import io
import os
import sys
import json
import asyncio
import traceback
from datetime import datetime

# 修正 Windows 終端編碼
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# 確保能 import backend 模組
backend_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'backend')
sys.path.insert(0, backend_dir)
os.environ.setdefault("GEMINI_API_KEY", "test-key-for-unit-tests")

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
from app.services.form_fill import FormFillService


# ================================================================
# 測試用 mock 資料
# ================================================================

MOCK_FILL_DATA = {
    "設備名稱": "三相感應馬達 A-01",
    "設備編號": "EQ-2026-0415",
    "檢查日期": "2026-04-15",
    "檢查人員": "張大明",
    "設備位置": "B棟 3F 配電室",
    "廠區": "南科廠區",
    "所屬廠區": "南科廠區",
    "溫度讀數": "62.5",
    "溫度": "62.5",
    "表面溫度": "62.5",
    "電壓值": "382",
    "電壓": "382",
    "轉速": "1480",
    "振動值": "3.2",
    "振動": "3.2",
    "整體狀態": "正常",
    "備註": "例行檢查，無異常",
    "絕緣電阻": "5.2",
    "接地電阻": "3.8",
    "A相電流": "12.3",
    "B相電流": "12.1",
    "C相電流": "12.5",
    "漏電流": "15",
    "噪音": "72",
    "電流": "12.3",
    # checkbox 值
    "是否合格": "true",
    "判定": "合格",
    "合格": "true",
    # L5 英文欄位
    "Equip. Name": "Motor A-01",
    "S/N": "SN-2026-0415",
    "Insp. Date": "2026-04-15",
    "Inspector": "David Chang",
    "Location": "Bldg-B 3F",
    "Dept": "Maintenance",
    "Ins. Res.": "5.2",
    "Gnd. Res.": "3.8",
    "Temp.": "62.5",
    "Vib.": "3.2",
    "RPM": "1480",
    # 清單項目
    "絕緣測試": "true",
    "接地測試": "true",
    "漏電斷路器動作": "true",
    "線路絕緣狀態": "true",
    "電氣接點狀態": "true",
    "防爆設備完整性": "false",
    "過載保護裝置": "true",
    "緊急停止裝置": "true",
    "漏電斷路器": "true",
    "漏電保護": "true",
    "軸承溫度": "58.3",
    "皮帶張力": "正常",
    "粉塵量": "6.2",
    "軸承狀態": "true",
    "絕緣電阻(MΩ)": "5.2",
    "接地電阻(Ω)": "3.8",
    "A相電流(A)": "12.3",
    "B相電流(A)": "12.1",
    "C相電流(A)": "12.5",
    "漏電流(mA)": "15",
    "電壓(V)": "382",
    "溫度(℃)": "62.5",
    "噪音(dB)": "72",
    "振動(mm/s)": "3.2",
    "表面溫度(℃)": "62.5",
    "電壓(V)：": "382",
    "電流(A)：": "12.3",
    "溫度(℃)：": "62.5",
    "振動(mm/s)：": "3.2",
    "轉速(RPM)": "1480",
    # L3 混合表
    "檢查人員簽章": "張大明",
    "主管簽章": "李主管",
    "簽核日期": "2026-04-15",
}


# ================================================================
# 讀取 Excel 所有格內容（before/after 比對用）
# ================================================================

def read_excel_cells(file_bytes: bytes) -> dict:
    """讀取 Excel 所有 sheet 的所有有值格，回傳 {sheet_name: {cell_ref: value}}"""
    wb = load_workbook(io.BytesIO(file_bytes))
    result = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        cells = {}
        max_row = min(ws.max_row or 1, 300)
        max_col = min(ws.max_column or 1, 50)
        for r in range(1, max_row + 1):
            for c in range(1, max_col + 1):
                val = ws.cell(row=r, column=c).value
                if val is not None and str(val).strip():
                    ref = f"{get_column_letter(c)}{r}"
                    cells[ref] = str(val)
        result[sheet_name] = cells
    return result


def read_word_content(file_bytes: bytes) -> dict:
    """讀取 Word 段落和表格內容"""
    doc = Document(io.BytesIO(file_bytes))
    result = {"paragraphs": [], "tables": []}

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            result["paragraphs"].append({"index": i, "text": para.text})

    for ti, table in enumerate(doc.tables):
        table_data = []
        for ri, row in enumerate(table.rows):
            row_data = []
            for ci, cell in enumerate(row.cells):
                row_data.append(cell.text)
            table_data.append(row_data)
        result["tables"].append({"table_index": ti, "rows": table_data})

    return result


# ================================================================
# 單一測試執行器
# ================================================================

async def run_single_test(service: FormFillService, file_path: str) -> dict:
    """
    對單一檔案執行完整測試流程。
    回傳 test result dict。
    """
    file_name = os.path.basename(file_path)
    file_ext = file_name.split('.')[-1].lower()

    print(f"  📄 {file_name} ...", end=" ", flush=True)

    result = {
        "file_name": file_name,
        "file_ext": file_ext,
        "status": "pending",
        "errors": [],
        "warnings": [],
        "structure_analysis": None,
        "fill_results": [],
        "before_content": None,
        "after_content": None,
        "field_count": 0,
        "filled_count": 0,
        "skipped_count": 0,
    }

    try:
        # 讀取檔案
        with open(file_path, 'rb') as f:
            file_content = f.read()

        # ---- BEFORE: 讀取填入前的內容 ----
        if file_ext == 'xlsx':
            result["before_content"] = read_excel_cells(file_content)
        elif file_ext == 'docx':
            result["before_content"] = read_word_content(file_content)

        # ---- Step 1: 結構分析 ----
        try:
            analysis = await service.analyze_structure(file_content, file_name)
        except Exception as e:
            result["status"] = "FAIL"
            result["errors"].append(f"結構分析失敗: {e}")
            print("❌ 結構分析失敗")
            return result

        result["structure_analysis"] = analysis
        field_map = analysis.get("field_map", [])
        result["field_count"] = len(field_map)

        if not field_map:
            result["status"] = "WARN"
            result["warnings"].append("結構分析未偵測到任何欄位")
            print(f"⚠️  0 欄位")
            return result

        # ---- Step 2: 建立 fill_values（模擬 AI 映射） ----
        fill_values = []
        matched_fields = []
        unmatched_fields = []

        for field in field_map:
            fid = field["field_id"]
            fname = field["field_name"].rstrip("：: ")
            ftype = field.get("field_type", "text")
            val_loc = field.get("value_location")

            # 嘗試從 MOCK_FILL_DATA 配對
            matched_value = None
            for key, val in MOCK_FILL_DATA.items():
                # 完全匹配或部分匹配
                if key == fname or key in fname or fname in key:
                    matched_value = val
                    break

            if matched_value and val_loc:
                fill_values.append({
                    "field_id": fid,
                    "value": matched_value,
                    "confidence": 0.92,
                    "source": "mock_data",
                })
                matched_fields.append(fname)
            else:
                unmatched_fields.append(fname)

        result["filled_count"] = len(fill_values)
        result["skipped_count"] = len(unmatched_fields)

        if unmatched_fields:
            result["warnings"].append(
                f"以下欄位無對應 mock 資料: {', '.join(unmatched_fields[:10])}"
            )

        # ---- Step 3: 預覽 ----
        try:
            preview = await service.preview_auto_fill(field_map, fill_values)
            result["preview_warnings"] = preview.get("warnings", [])
        except Exception as e:
            result["warnings"].append(f"預覽失敗: {e}")

        # ---- Step 4: 執行回填 ----
        if fill_values:
            try:
                filled_bytes = await service.auto_fill(
                    file_content, file_name, field_map, fill_values
                )

                # ---- AFTER: 讀取填入後的內容 ----
                if file_ext == 'xlsx':
                    result["after_content"] = read_excel_cells(filled_bytes)
                elif file_ext == 'docx':
                    result["after_content"] = read_word_content(filled_bytes)

                # ---- Step 5: 驗證回填結果 ----
                fill_results = []
                for fv in fill_values:
                    fid = fv["field_id"]
                    field = next((f for f in field_map if f["field_id"] == fid), None)
                    if not field:
                        continue

                    val_loc = field.get("value_location", {})
                    expected = fv["value"]
                    actual = None
                    cell_ref = ""

                    if file_ext == 'xlsx' and val_loc:
                        sheet = val_loc.get("sheet", "")
                        cell_ref = val_loc.get("cell", "")
                        after_sheet = (result["after_content"] or {}).get(sheet, {})
                        actual = after_sheet.get(cell_ref, "(空)")
                        before_sheet = (result["before_content"] or {}).get(sheet, {})
                        before_val = before_sheet.get(cell_ref, "(空)")
                    elif file_ext == 'docx' and val_loc:
                        loc_type = val_loc.get("type", "")
                        cell_ref = f"{loc_type}[{val_loc.get('paragraph_index', val_loc.get('table_index', '?'))}]"
                        before_val = "(原始)"
                        actual = "(已回填)"
                        # Word 的驗證比較複雜，讀 after content 的段落/表格
                        if loc_type == "paragraph":
                            pi = val_loc.get("paragraph_index", -1)
                            after_paras = (result.get("after_content") or {}).get("paragraphs", [])
                            matched_para = next((p for p in after_paras if p["index"] == pi), None)
                            if matched_para:
                                actual = matched_para["text"]
                            before_paras = (result.get("before_content") or {}).get("paragraphs", [])
                            matched_bp = next((p for p in before_paras if p["index"] == pi), None)
                            if matched_bp:
                                before_val = matched_bp["text"]
                        elif loc_type == "table":
                            ti = val_loc.get("table_index", -1)
                            ri = val_loc.get("row_index", -1)
                            ci = val_loc.get("cell_index", -1)
                            after_tables = (result.get("after_content") or {}).get("tables", [])
                            if ti < len(after_tables):
                                rows = after_tables[ti].get("rows", [])
                                if ri < len(rows) and ci < len(rows[ri]):
                                    actual = rows[ri][ci]
                            before_tables = (result.get("before_content") or {}).get("tables", [])
                            if ti < len(before_tables):
                                rows = before_tables[ti].get("rows", [])
                                if ri < len(rows) and ci < len(rows[ri]):
                                    before_val = rows[ri][ci]

                    # 判定回填是否成功
                    # 數值型：比較數值相等
                    # 文字型：包含即可（因為 checkbox 會轉換）
                    success = False
                    if actual and actual != "(空)":
                        if str(expected) in str(actual) or str(actual) in str(expected):
                            success = True
                        # 特別處理 checkbox 轉換
                        if expected in ("true", "是", "合格", "正常") and actual in ("合格", "✓", "☑"):
                            success = True
                        if expected in ("false", "否", "不合格", "異常") and actual in ("不合格", "✗", "☒"):
                            success = True
                        # 數值比較
                        try:
                            if float(expected) == float(actual):
                                success = True
                        except (ValueError, TypeError):
                            pass

                    fill_results.append({
                        "field_name": field["field_name"],
                        "field_type": field.get("field_type", "text"),
                        "location": cell_ref,
                        "before": before_val if 'before_val' in dir() else "(未知)",
                        "expected": expected,
                        "actual": actual or "(空)",
                        "success": success,
                    })

                result["fill_results"] = fill_results
                passed = sum(1 for r in fill_results if r["success"])
                failed = sum(1 for r in fill_results if not r["success"])

                if failed == 0:
                    result["status"] = "PASS"
                    print(f"✅ {result['field_count']} 欄位, {passed} 填入成功")
                else:
                    result["status"] = "PARTIAL"
                    print(f"⚠️  {passed}/{passed+failed} 成功, {failed} 失敗")

            except Exception as e:
                result["status"] = "FAIL"
                result["errors"].append(f"回填執行失敗: {e}\n{traceback.format_exc()}")
                print(f"❌ 回填失敗: {e}")
        else:
            result["status"] = "WARN"
            result["warnings"].append("無可填入的值（全部欄位都無 mock 資料配對）")
            print(f"⚠️  {result['field_count']} 欄位偵測到，但 0 配對")

    except Exception as e:
        result["status"] = "FAIL"
        result["errors"].append(f"測試執行錯誤: {e}\n{traceback.format_exc()}")
        print(f"❌ {e}")

    return result


# ================================================================
# HTML 報告產生
# ================================================================

def generate_html_report(all_results: dict, output_path: str):
    """產生完整 HTML 測試報告"""

    # 統計
    total_files = sum(len(tests) for tests in all_results.values())
    total_pass = sum(
        1 for tests in all_results.values()
        for t in tests if t["status"] == "PASS"
    )
    total_partial = sum(
        1 for tests in all_results.values()
        for t in tests if t["status"] == "PARTIAL"
    )
    total_fail = sum(
        1 for tests in all_results.values()
        for t in tests if t["status"] == "FAIL"
    )
    total_warn = sum(
        1 for tests in all_results.values()
        for t in tests if t["status"] == "WARN"
    )

    status_emoji = {"PASS": "✅", "PARTIAL": "⚠️", "FAIL": "❌", "WARN": "🟡", "pending": "⏳"}
    status_color = {"PASS": "#27ae60", "PARTIAL": "#f39c12", "FAIL": "#e74c3c", "WARN": "#f1c40f", "pending": "#95a5a6"}

    html_parts = []
    html_parts.append(f"""<!DOCTYPE html>
<html lang="zh-TW">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>InduSpect 表單自動回填測試報告</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: 'Segoe UI', 'Microsoft JhengHei', sans-serif; background: #f5f6fa; color: #2c3e50; }}
.container {{ max-width: 1200px; margin: 0 auto; padding: 20px; }}
h1 {{ text-align: center; margin: 30px 0; color: #2c3e50; }}
.summary {{ display: flex; gap: 20px; justify-content: center; margin: 30px 0; flex-wrap: wrap; }}
.summary-card {{ background: white; border-radius: 12px; padding: 20px 30px; box-shadow: 0 2px 10px rgba(0,0,0,0.08); text-align: center; min-width: 140px; }}
.summary-card .number {{ font-size: 36px; font-weight: bold; }}
.summary-card .label {{ font-size: 14px; color: #7f8c8d; margin-top: 5px; }}
.level-section {{ background: white; border-radius: 12px; margin: 25px 0; box-shadow: 0 2px 10px rgba(0,0,0,0.08); overflow: hidden; }}
.level-header {{ padding: 18px 25px; font-size: 18px; font-weight: bold; cursor: pointer; display: flex; justify-content: space-between; align-items: center; }}
.level-header:hover {{ opacity: 0.9; }}
.level-body {{ padding: 0 25px 25px; }}
.test-file {{ border: 1px solid #ecf0f1; border-radius: 8px; margin: 15px 0; overflow: hidden; }}
.test-file-header {{ padding: 12px 18px; background: #f8f9fa; display: flex; justify-content: space-between; align-items: center; font-weight: 600; cursor: pointer; }}
.test-file-header:hover {{ background: #ecf0f1; }}
.test-file-body {{ padding: 15px 18px; display: none; }}
.test-file-body.open {{ display: block; }}
.info-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin: 10px 0; }}
.info-item {{ padding: 8px 12px; background: #f8f9fa; border-radius: 6px; font-size: 14px; }}
.info-label {{ color: #7f8c8d; font-size: 12px; }}
table {{ width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 13px; }}
th {{ background: #34495e; color: white; padding: 10px 12px; text-align: left; position: sticky; top: 0; }}
td {{ padding: 8px 12px; border-bottom: 1px solid #ecf0f1; }}
tr:hover {{ background: #f8f9fa; }}
.badge {{ display: inline-block; padding: 2px 10px; border-radius: 12px; font-size: 12px; font-weight: 600; color: white; }}
.field-map-table {{ margin-top: 10px; }}
.field-map-table th {{ background: #2c3e50; font-size: 12px; }}
.field-map-table td {{ font-size: 12px; font-family: 'Consolas', monospace; }}
.diff-added {{ background: #d4edda; }}
.diff-empty {{ color: #bdc3c7; font-style: italic; }}
.error-box {{ background: #fce4e4; border-left: 4px solid #e74c3c; padding: 12px; margin: 10px 0; border-radius: 4px; font-size: 13px; }}
.warn-box {{ background: #fef9e7; border-left: 4px solid #f39c12; padding: 12px; margin: 10px 0; border-radius: 4px; font-size: 13px; }}
.section-title {{ font-size: 15px; font-weight: 600; margin: 20px 0 10px; padding-bottom: 6px; border-bottom: 2px solid #3498db; color: #2c3e50; }}
.timestamp {{ text-align: center; color: #95a5a6; font-size: 13px; margin-bottom: 30px; }}
details {{ margin: 5px 0; }}
summary {{ cursor: pointer; color: #3498db; font-size: 13px; }}
.fill-result-row td:nth-child(4), .fill-result-row td:nth-child(5) {{
    max-width: 200px; word-break: break-all;
}}
</style>
<script>
function toggleBody(el) {{
    const body = el.nextElementSibling;
    if (body) body.classList.toggle('open');
}}
function expandAll() {{
    document.querySelectorAll('.test-file-body').forEach(b => b.classList.add('open'));
}}
function collapseAll() {{
    document.querySelectorAll('.test-file-body').forEach(b => b.classList.remove('open'));
}}
</script>
</head>
<body>
<div class="container">
<h1>InduSpect 表單自動回填測試報告</h1>
<p class="timestamp">產生時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 測試檔案: {total_files} 個</p>

<div class="summary">
  <div class="summary-card"><div class="number" style="color:#27ae60">{total_pass}</div><div class="label">通過 PASS</div></div>
  <div class="summary-card"><div class="number" style="color:#f39c12">{total_partial}</div><div class="label">部分 PARTIAL</div></div>
  <div class="summary-card"><div class="number" style="color:#e74c3c">{total_fail}</div><div class="label">失敗 FAIL</div></div>
  <div class="summary-card"><div class="number" style="color:#f1c40f">{total_warn}</div><div class="label">警告 WARN</div></div>
  <div class="summary-card"><div class="number" style="color:#2c3e50">{total_files}</div><div class="label">總計</div></div>
</div>

<div style="text-align:center; margin: 15px 0;">
  <button onclick="expandAll()" style="padding:6px 16px; cursor:pointer; border:1px solid #ccc; border-radius:4px; margin:0 5px;">全部展開</button>
  <button onclick="collapseAll()" style="padding:6px 16px; cursor:pointer; border:1px solid #ccc; border-radius:4px; margin:0 5px;">全部收合</button>
</div>
""")

    level_colors = {
        "L1": "#27ae60",
        "L2": "#2ecc71",
        "L3": "#f39c12",
        "L4": "#e67e22",
        "L5": "#e74c3c",
    }
    level_labels = {
        "L1": "Level 1：簡單 Key-Value 表",
        "L2": "Level 2：標準清單表",
        "L3": "Level 3：混合結構表",
        "L4": "Level 4：合併儲存格 + 雙欄 Checkbox",
        "L5": "Level 5：自由格式 / 非標準表",
    }

    for level in ["L1", "L2", "L3", "L4", "L5"]:
        tests = all_results.get(level, [])
        if not tests:
            continue

        level_pass = sum(1 for t in tests if t["status"] == "PASS")
        level_total = len(tests)
        bg_color = level_colors.get(level, "#95a5a6")

        html_parts.append(f"""
<div class="level-section">
  <div class="level-header" style="background: {bg_color}; color: white;">
    <span>{level_labels.get(level, level)} ({level_pass}/{level_total} 通過)</span>
    <span>{'⭐' * int(level[1])}</span>
  </div>
  <div class="level-body">
""")

        for test in tests:
            s = test["status"]
            s_color = status_color.get(s, "#95a5a6")
            s_emoji = status_emoji.get(s, "?")

            html_parts.append(f"""
    <div class="test-file">
      <div class="test-file-header" onclick="toggleBody(this)">
        <span>{s_emoji} {test['file_name']}</span>
        <span>
          <span class="badge" style="background:{s_color}">{s}</span>
          &nbsp;{test['field_count']} 欄位 | {test['filled_count']} 填入
        </span>
      </div>
      <div class="test-file-body">
""")

            # 基本資訊
            html_parts.append(f"""
        <div class="info-grid">
          <div class="info-item"><span class="info-label">檔案格式</span><br>{test['file_ext'].upper()}</div>
          <div class="info-item"><span class="info-label">偵測欄位數</span><br>{test['field_count']}</div>
          <div class="info-item"><span class="info-label">成功填入數</span><br>{test['filled_count']}</div>
          <div class="info-item"><span class="info-label">未配對數</span><br>{test['skipped_count']}</div>
        </div>
""")

            # 錯誤訊息
            for err in test.get("errors", []):
                html_parts.append(f'<div class="error-box">❌ {err}</div>')

            # 警告訊息
            for warn in test.get("warnings", []):
                html_parts.append(f'<div class="warn-box">⚠️ {warn}</div>')

            # 結構分析結果 — 欄位地圖
            analysis = test.get("structure_analysis")
            if analysis and analysis.get("field_map"):
                html_parts.append('<div class="section-title">欄位偵測結果（Field Map）</div>')
                html_parts.append("""
        <table class="field-map-table">
          <tr><th>Field ID</th><th>欄位名稱</th><th>類型</th><th>標籤位置</th><th>值位置</th><th>方向</th></tr>
""")
                for fm in analysis["field_map"]:
                    label_loc = fm.get("label_location", {})
                    val_loc = fm.get("value_location", {})
                    label_pos = f"{label_loc.get('sheet','')}/{label_loc.get('cell','')}" if label_loc.get("cell") else f"{label_loc.get('type','')}"
                    val_pos = f"{val_loc.get('sheet','')}/{val_loc.get('cell','')}" if val_loc and val_loc.get("cell") else (f"{val_loc.get('type','')}" if val_loc else "—")
                    direction = val_loc.get("direction", "—") if val_loc else "—"

                    html_parts.append(f"""
          <tr>
            <td>{fm['field_id']}</td>
            <td><strong>{fm['field_name']}</strong></td>
            <td>{fm.get('field_type','text')}</td>
            <td>{label_pos}</td>
            <td>{val_pos}</td>
            <td>{direction}</td>
          </tr>""")
                html_parts.append("</table>")

            # 回填結果 — Before / After 對照表
            fill_results = test.get("fill_results", [])
            if fill_results:
                passed_count = sum(1 for r in fill_results if r["success"])
                failed_count = sum(1 for r in fill_results if not r["success"])

                html_parts.append(f"""
        <div class="section-title">回填結果對照（{passed_count} 成功 / {failed_count} 失敗）</div>
        <table>
          <tr>
            <th>欄位名稱</th>
            <th>類型</th>
            <th>位置</th>
            <th>Before（填入前）</th>
            <th>After（填入後）</th>
            <th>預期值</th>
            <th>結果</th>
          </tr>
""")
                for fr in fill_results:
                    row_class = "diff-added" if fr["success"] else ""
                    before_class = "diff-empty" if not fr["before"] or fr["before"] in ("(空)", "", "______", "____", "☐") else ""
                    result_badge = f'<span class="badge" style="background:#27ae60">✓ 通過</span>' if fr["success"] else f'<span class="badge" style="background:#e74c3c">✗ 失敗</span>'

                    # 顯示 before 值
                    before_display = fr.get("before", "(空)")
                    if before_display in ("", None):
                        before_display = "(空)"

                    html_parts.append(f"""
          <tr class="fill-result-row {row_class}">
            <td><strong>{fr['field_name']}</strong></td>
            <td>{fr['field_type']}</td>
            <td><code>{fr['location']}</code></td>
            <td class="{before_class}">{before_display}</td>
            <td><strong>{fr['actual']}</strong></td>
            <td>{fr['expected']}</td>
            <td>{result_badge}</td>
          </tr>""")
                html_parts.append("</table>")

            html_parts.append("</div></div>")  # test-file-body, test-file

        html_parts.append("</div></div>")  # level-body, level-section

    html_parts.append("""
</div>
</body>
</html>""")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_parts))

    print(f"\n📊 報告已產生: {output_path}")


# ================================================================
# 主程式
# ================================================================

async def main():
    print("=" * 60)
    print("InduSpect 表單自動回填分級測試")
    print("=" * 60)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    service = FormFillService()

    # 定義各 Level 的測試檔案
    test_files = {
        "L1": [
            "L1_simple_kv.xlsx",
            "L1_simple_kv.docx",
        ],
        "L2": [
            "L2_standard_list.xlsx",
            "L2_list_with_sections.xlsx",
        ],
        "L3": [
            "L3_mixed_structure.xlsx",
            "L3_mixed_structure.docx",
        ],
        "L4": [
            "L4_merged_cells.xlsx",
            "L4_dual_checkbox.xlsx",
            "L4_merged_checkbox_combo.xlsx",
        ],
        "L5": [
            "L5_nonstandard_labels.xlsx",
            "L5_oversized.xlsx",
            "L5_multi_sheet.xlsx",
            "L5_nested_table.docx",
            "L5_irregular_merge.xlsx",
        ],
    }

    all_results = {}

    for level, files in test_files.items():
        print(f"\n{'─' * 40}")
        print(f"🔷 {level}")
        print(f"{'─' * 40}")

        level_results = []
        for fname in files:
            fpath = os.path.join(script_dir, fname)
            if not os.path.exists(fpath):
                print(f"  📄 {fname} ... ❌ 檔案不存在")
                level_results.append({
                    "file_name": fname,
                    "file_ext": fname.split('.')[-1],
                    "status": "FAIL",
                    "errors": [f"檔案不存在: {fpath}"],
                    "warnings": [],
                    "structure_analysis": None,
                    "fill_results": [],
                    "before_content": None,
                    "after_content": None,
                    "field_count": 0,
                    "filled_count": 0,
                    "skipped_count": 0,
                })
                continue

            result = await run_single_test(service, fpath)
            level_results.append(result)

        all_results[level] = level_results

    # 產生 HTML 報告
    report_path = os.path.join(script_dir, "AUTO_FILL_TEST_REPORT.html")
    generate_html_report(all_results, report_path)

    # 摘要
    print(f"\n{'=' * 60}")
    print("測試摘要")
    print(f"{'=' * 60}")
    for level, tests in all_results.items():
        pass_count = sum(1 for t in tests if t["status"] == "PASS")
        total = len(tests)
        status_str = "✅ ALL PASS" if pass_count == total else f"⚠️  {pass_count}/{total}"
        print(f"  {level}: {status_str}")
    print()


if __name__ == "__main__":
    asyncio.run(main())
