"""
照片處理服務 - 照片自動插入報告

從 form_fill.py 提取的照片相關方法
"""

import io
import logging
import base64
from typing import Optional

from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XlImage
from openpyxl.styles import Font, Alignment, Border, Side
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage

logger = logging.getLogger(__name__)


class PhotoProcessingService:
    """照片處理服務 - 負責將照片插入 Excel/Word 報告"""

    async def insert_photos_into_report(
        self,
        file_content: bytes,
        file_name: str,
        photo_bindings: list[dict],
    ) -> bytes:
        """
        將照片自動插入到 Excel/Word 報告中

        photo_bindings 格式:
        [
            {
                "task_id": "task_1",
                "display_name": "絕緣電阻測量",
                "photo_base64": "base64...",       # 照片 base64
                "photo_bytes": b"...",             # 或直接提供 bytes
                "capture_time": "2026-03-13 14:30",
                "sequence": 1,
            },
            ...
        ]
        """
        ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''

        if ext == 'xlsx':
            return await self._insert_photos_excel(file_content, photo_bindings)
        elif ext == 'docx':
            return await self._insert_photos_word(file_content, photo_bindings)
        else:
            raise ValueError(f"不支援的檔案類型: {ext}")

    async def _insert_photos_excel(
        self,
        file_content: bytes,
        photo_bindings: list[dict],
    ) -> bytes:
        """
        Excel 照片插入

        策略: 在最後新增一個「照片附件」工作表，包含：
        - 編號、檢查項目、現場照片、拍攝時間
        """
        wb = load_workbook(io.BytesIO(file_content))

        # 建立照片附件工作表
        sheet_name = "照片附件"
        idx = 1
        while sheet_name in wb.sheetnames:
            sheet_name = f"照片附件({idx})"
            idx += 1

        ws = wb.create_sheet(title=sheet_name)

        # 表頭樣式
        header_font = Font(bold=True, size=12)
        header_align = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin'),
        )

        # 標題行
        ws.merge_cells('A1:D1')
        title_cell = ws['A1']
        title_cell.value = "現場照片記錄"
        title_cell.font = Font(bold=True, size=14)
        title_cell.alignment = Alignment(horizontal='center', vertical='center')

        # 表頭
        headers = ["編號", "檢查項目", "現場照片", "拍攝時間"]
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 45
        ws.column_dimensions['D'].width = 20

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=2, column=col_idx, value=header)
            cell.font = header_font
            cell.alignment = header_align
            cell.border = thin_border

        # 資料列
        current_row = 3
        sorted_bindings = sorted(photo_bindings, key=lambda x: x.get("sequence", 0))

        for binding in sorted_bindings:
            photo_io = self._prepare_photo_for_insert(binding)
            if photo_io is None:
                continue

            display_name = binding.get("display_name", "未命名")
            capture_time = binding.get("capture_time", "")
            sequence = binding.get("sequence", current_row - 2)

            # 設定列高以容納照片（約 225px ≈ 170pt）
            ws.row_dimensions[current_row].height = 170

            # 編號
            cell_num = ws.cell(row=current_row, column=1, value=sequence)
            cell_num.alignment = Alignment(horizontal='center', vertical='center')
            cell_num.border = thin_border

            # 檢查項目
            cell_name = ws.cell(row=current_row, column=2, value=display_name)
            cell_name.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell_name.border = thin_border

            # 照片
            cell_photo = ws.cell(row=current_row, column=3, value="")
            cell_photo.border = thin_border

            img = XlImage(photo_io)
            img.width = 300
            img.height = 225
            anchor = f"C{current_row}"
            ws.add_image(img, anchor)

            # 拍攝時間
            cell_time = ws.cell(row=current_row, column=4, value=capture_time)
            cell_time.alignment = Alignment(horizontal='center', vertical='center')
            cell_time.border = thin_border

            current_row += 1

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output.read()

    async def _insert_photos_word(
        self,
        file_content: bytes,
        photo_bindings: list[dict],
    ) -> bytes:
        """
        Word 照片插入

        策略: 在文件末尾新增「照片記錄」章節
        """
        doc = Document(io.BytesIO(file_content))

        # 新增分頁和標題
        doc.add_page_break()
        doc.add_heading("照片記錄", level=1)

        sorted_bindings = sorted(photo_bindings, key=lambda x: x.get("sequence", 0))

        for binding in sorted_bindings:
            photo_io = self._prepare_photo_for_insert(binding)
            if photo_io is None:
                continue

            display_name = binding.get("display_name", "未命名")
            capture_time = binding.get("capture_time", "")
            sequence = binding.get("sequence", 0)

            # 項目標題
            p_title = doc.add_paragraph()
            run = p_title.add_run(f"{sequence}. {display_name}")
            run.bold = True
            run.font.size = Pt(12)

            # 插入照片（寬度 12cm）
            p_photo = doc.add_paragraph()
            p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_photo.add_run()
            run_img.add_picture(photo_io, width=Cm(12))

            # 拍攝時間
            if capture_time:
                p_time = doc.add_paragraph()
                p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_time = p_time.add_run(f"拍攝時間: {capture_time}")
                run_time.font.size = Pt(9)
                run_time.font.color.rgb = None  # 預設色

            # 間距
            doc.add_paragraph("")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _prepare_photo_for_insert(
        self,
        binding: dict,
        max_width_px: int = 600,
        max_height_px: int = 450,
        max_size_kb: int = 500,
    ) -> Optional[io.BytesIO]:
        """
        準備照片用於插入文件

        1. 從 base64 或 bytes 解碼
        2. 縮放到合理大小
        3. 壓縮至 max_size_kb 以下
        4. 回傳 BytesIO
        """
        try:
            # 取得照片原始資料
            photo_bytes = binding.get("photo_bytes")
            photo_base64 = binding.get("photo_base64")

            if photo_bytes:
                if isinstance(photo_bytes, str):
                    photo_bytes = photo_bytes.encode('latin-1')
                img_data = photo_bytes
            elif photo_base64:
                # 處理可能有 data:image/...;base64, 前綴的情況
                if ',' in photo_base64:
                    photo_base64 = photo_base64.split(',', 1)[1]
                img_data = base64.b64decode(photo_base64)
            else:
                logger.warning(f"Binding {binding.get('task_id')} has no photo data")
                return None

            img = PILImage.open(io.BytesIO(img_data))

            # 轉為 RGB（處理 RGBA 或 P 模式）
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')

            # 按比例縮放
            w, h = img.size
            if w > max_width_px or h > max_height_px:
                ratio = min(max_width_px / w, max_height_px / h)
                new_w = int(w * ratio)
                new_h = int(h * ratio)
                img = img.resize((new_w, new_h), PILImage.LANCZOS)

            # 壓縮：嘗試不同品質直到低於目標大小
            quality = 85
            while quality >= 30:
                output = io.BytesIO()
                img.save(output, format='JPEG', quality=quality, optimize=True)
                size_kb = output.tell() / 1024
                if size_kb <= max_size_kb:
                    output.seek(0)
                    return output
                quality -= 10

            # 即使最低品質仍超過限制，還是回傳
            output.seek(0)
            return output

        except Exception as e:
            logger.error(f"Prepare photo failed for {binding.get('task_id')}: {e}")
            return None
