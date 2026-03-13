"""
自動回填服務 — 執行回填、預覽、報告產生

職責：
- 執行 Excel/Word 自動回填（保留原始格式）
- 預覽回填結果
- 舊版報告產生相容
"""

import io
import re
import copy
import json
import logging
from typing import Optional
from datetime import datetime

import google.generativeai as genai
from openpyxl import load_workbook
from docx import Document

from app.config import settings
from app.services.form_utils import (
    convert_value, replace_paragraph_text_preserve_format,
    guess_field_type,
)

logger = logging.getLogger(__name__)


class AutoFillService:
    """自動回填服務"""

    def __init__(self):
        genai.configure(api_key=settings.gemini_api_key)

    # ================================================================
    # 自動回填引擎
    # ================================================================

    async def auto_fill(
        self,
        file_content: bytes,
        file_name: str,
        field_map: list[dict],
        fill_values: list[dict],
    ) -> bytes:
        """
        執行自動回填：將值寫入原始文件的指定位置

        Args:
            file_content: 原始文件內容
            file_name: 文件名稱（判斷格式用）
            field_map: 欄位位置地圖（含 value_location）
            fill_values: 要填入的值列表
                [{"field_id": "...", "value": "..."}, ...]

        Returns:
            回填後的文件 bytes
        """
        file_type = file_name.split('.')[-1].lower()

        # 建立 field_id -> value 的查找表
        value_lookup = {fv["field_id"]: fv["value"] for fv in fill_values}

        # 建立 field_id -> field_map entry 的查找表
        field_lookup = {f["field_id"]: f for f in field_map}

        if file_type == 'xlsx':
            return await self._auto_fill_excel(
                file_content, field_lookup, value_lookup
            )
        elif file_type == 'docx':
            return await self._auto_fill_word(
                file_content, field_lookup, value_lookup
            )
        else:
            raise ValueError(f"不支援的檔案格式: {file_type}")

    async def _auto_fill_excel(
        self,
        file_content: bytes,
        field_lookup: dict,
        value_lookup: dict,
    ) -> bytes:
        """
        回填 Excel 檔案

        保留原始格式（字體、邊框、合併儲存格、樣式等）
        """
        wb = load_workbook(io.BytesIO(file_content))

        for field_id, value in value_lookup.items():
            field = field_lookup.get(field_id)
            if not field:
                continue

            val_loc = field.get("value_location")
            if not val_loc:
                # 沒有找到值儲存格，嘗試寫入標籤儲存格
                label_loc = field.get("label_location", {})
                sheet_name = label_loc.get("sheet")
                cell_coord = label_loc.get("cell")
                if sheet_name and cell_coord and sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    ws[cell_coord] = value
                continue

            sheet_name = val_loc.get("sheet")
            cell_coord = val_loc.get("cell")

            if not sheet_name or not cell_coord:
                continue

            if sheet_name not in wb.sheetnames:
                logger.warning(f"Sheet '{sheet_name}' not found, skipping field {field_id}")
                continue

            ws = wb[sheet_name]

            # 根據欄位類型轉換值
            typed_value = convert_value(value, field.get("field_type", "text"))

            # 寫入值，保留原始格式
            target_cell = ws[cell_coord]

            # 複製原始格式資訊
            original_font = copy.copy(target_cell.font) if target_cell.font else None
            original_alignment = copy.copy(target_cell.alignment) if target_cell.alignment else None
            original_number_format = target_cell.number_format

            target_cell.value = typed_value

            # 還原格式
            if original_font:
                target_cell.font = original_font
            if original_alignment:
                target_cell.alignment = original_alignment
            if original_number_format:
                target_cell.number_format = original_number_format

        # 輸出為 bytes
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output.read()

    async def _auto_fill_word(
        self,
        file_content: bytes,
        field_lookup: dict,
        value_lookup: dict,
    ) -> bytes:
        """
        回填 Word 檔案

        保留原始格式（字體、段落樣式等）
        """
        doc = Document(io.BytesIO(file_content))

        for field_id, value in value_lookup.items():
            field = field_lookup.get(field_id)
            if not field:
                continue

            val_loc = field.get("value_location")
            if not val_loc:
                continue

            loc_type = val_loc.get("type")

            if loc_type == "paragraph":
                para_idx = val_loc.get("paragraph_index")
                if para_idx is not None and para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    replace_pattern = val_loc.get("replace_pattern", "after_colon")

                    if replace_pattern == "after_colon":
                        # 替換冒號/：後面的內容
                        text = para.text
                        for sep in ['：', ':']:
                            if sep in text:
                                prefix = text.split(sep)[0] + sep
                                replace_paragraph_text_preserve_format(
                                    para, f"{prefix} {value}"
                                )
                                break
                    else:
                        # 替換佔位符
                        replace_paragraph_text_preserve_format(para, value)

            elif loc_type == "table":
                table_idx = val_loc.get("table_index")
                row_idx = val_loc.get("row_index")
                cell_idx = val_loc.get("cell_index")

                if (table_idx is not None and
                        table_idx < len(doc.tables) and
                        row_idx is not None and
                        cell_idx is not None):
                    table = doc.tables[table_idx]
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        if cell_idx < len(row.cells):
                            cell = row.cells[cell_idx]
                            # 保留格式寫入
                            if cell.paragraphs:
                                replace_paragraph_text_preserve_format(
                                    cell.paragraphs[0], str(value)
                                )
                            else:
                                cell.text = str(value)

        # 輸出為 bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    # ================================================================
    # 預覽回填
    # ================================================================

    async def preview_fill(
        self,
        template: dict,
        inspection_data: dict
    ) -> dict:
        """預覽填入結果（舊版 API 相容）"""
        field_values = {}
        warnings = []

        for field in template.get("fields", []):
            mapping = field.get("mapping")
            if mapping and mapping in inspection_data:
                value = inspection_data[mapping]
                if isinstance(value, dict):
                    value = str(value)
                field_values[field["field_name"]] = str(value) if value else ""
            else:
                field_values[field["field_name"]] = ""
                if mapping:
                    warnings.append(f"欄位 '{field['field_name']}' 對應的資料不存在")

        return {
            "template_name": template.get("name", ""),
            "vendor_name": template.get("vendor_name", ""),
            "field_values": field_values,
            "warnings": warnings,
        }

    async def preview_auto_fill(
        self,
        field_map: list[dict],
        fill_values: list[dict],
    ) -> dict:
        """
        預覽自動回填結果（新版 API）

        回傳每個欄位即將填入的值以及信心度標記
        """
        value_lookup = {fv["field_id"]: fv for fv in fill_values}

        preview_items = []
        warnings = []

        for field in field_map:
            fid = field["field_id"]
            fv = value_lookup.get(fid)

            item = {
                "field_id": fid,
                "field_name": field["field_name"],
                "field_type": field.get("field_type", "text"),
                "value": fv["value"] if fv else None,
                "confidence": fv.get("confidence", 0.0) if fv else 0.0,
                "source": fv.get("source", "") if fv else "",
                "has_target": field.get("value_location") is not None,
            }
            preview_items.append(item)

            if not fv:
                warnings.append(f"欄位 '{field['field_name']}' 無對應值")
            elif item["confidence"] < 0.7:
                warnings.append(
                    f"欄位 '{field['field_name']}' 映射信心度較低 ({item['confidence']:.0%})，建議確認"
                )
            if not item["has_target"]:
                warnings.append(
                    f"欄位 '{field['field_name']}' 找不到值儲存格位置，無法回填"
                )

        return {
            "preview_items": preview_items,
            "total_fields": len(field_map),
            "filled_count": sum(1 for p in preview_items if p["value"] is not None),
            "warnings": warnings,
        }

    # ================================================================
    # 報告生成（舊版 API 相容）
    # ================================================================

    async def generate_report(
        self,
        report_id: str,
        template: dict,
        inspection_data: dict,
        output_format: str = "xlsx"
    ) -> str:
        """產生報告，回傳 output_path"""
        try:
            if template["file_type"] == "xlsx":
                output_path = await self._fill_excel(template, inspection_data, report_id)
            elif template["file_type"] == "docx":
                output_path = await self._fill_word(template, inspection_data, report_id)
            else:
                raise ValueError(f"Unsupported template type: {template['file_type']}")

            return output_path

        except Exception as e:
            logger.error(f"Generate report failed: {e}")
            raise

    async def _fill_excel(
        self,
        template: dict,
        inspection_data: dict,
        report_id: str
    ) -> str:
        """填入 Excel 模板"""
        wb = load_workbook(io.BytesIO(template["file_content"]))
        ws = wb.active

        for field in template["fields"]:
            mapping = field.get("mapping")
            if mapping and mapping in inspection_data:
                value = inspection_data[mapping]
                location = field["location"]
                ws[location] = value

        output_path = f"/tmp/report_{report_id}.xlsx"
        wb.save(output_path)

        return output_path

    async def _fill_word(
        self,
        template: dict,
        inspection_data: dict,
        report_id: str
    ) -> str:
        """填入 Word 模板"""
        doc = Document(io.BytesIO(template["file_content"]))

        for para in doc.paragraphs:
            for field in template["fields"]:
                mapping = field.get("mapping")
                if mapping and mapping in inspection_data:
                    value = str(inspection_data[mapping] or "")
                    placeholder = f"{{{{{field['field_name']}}}}}"
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, value)

        output_path = f"/tmp/report_{report_id}.docx"
        doc.save(output_path)

        return output_path
