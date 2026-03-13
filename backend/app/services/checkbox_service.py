"""
勾選式表格偵測與回填服務

從 form_fill.py 抽取的 checkbox 相關功能，
包含勾選雙欄偵測、勾選回填引擎等。
"""

import io
import re
import copy
import logging
from typing import Optional

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from docx import Document

logger = logging.getLogger(__name__)


class CheckboxService:
    """勾選式表格偵測與回填服務"""

    # ================================================================
    # 勾選式表格偵測 (Sprint 2 Task 2.1 新增)
    # ================================================================

    async def detect_checkbox_columns(
        self,
        file_content: bytes,
        file_name: str,
    ) -> dict:
        """
        偵測表單中的「合格/不合格」雙欄勾選結構

        掃描 Excel 表頭列，找到含有「合格」「不合格」「正常」「異常」等字樣的雙欄結構。
        回傳 dual_column_fields 清單，每項包含 pass_cell / fail_cell / remarks_cell。
        """
        file_type = file_name.split('.')[-1].lower()

        if file_type == 'xlsx':
            return await self._detect_checkbox_columns_excel(file_content)
        elif file_type == 'docx':
            return await self._detect_checkbox_columns_word(file_content)
        else:
            return {"dual_column_fields": [], "check_symbol": "✓"}

    async def _detect_checkbox_columns_excel(self, content: bytes) -> dict:
        """偵測 Excel 中的勾選雙欄結構"""
        wb = load_workbook(io.BytesIO(content))
        dual_fields = []
        detected_symbol = "✓"

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            max_row = min(ws.max_row or 1, 200)
            max_col = min(ws.max_column or 1, 50)

            # Phase 1: 找表頭列（含「合格/不合格」「正常/異常」等配對的列）
            header_rows = self._find_checkbox_header_rows(ws, max_row, max_col)

            for header_info in header_rows:
                header_row = header_info["row"]
                pass_col = header_info["pass_col"]
                fail_col = header_info["fail_col"]
                remarks_col = header_info.get("remarks_col")
                item_col = header_info.get("item_col", 1)

                # Phase 2: 掃描表頭下方的資料列
                for data_row in range(header_row + 1, max_row + 1):
                    # 取得檢查項目名稱
                    item_cell = ws.cell(row=data_row, column=item_col)
                    item_name = str(item_cell.value).strip() if item_cell.value else ""

                    if not item_name:
                        continue

                    # 跳過子表頭或空列
                    if self._is_section_header(item_name) or self._is_non_field_item(item_name):
                        continue

                    pass_coord = f"{get_column_letter(pass_col)}{data_row}"
                    fail_coord = f"{get_column_letter(fail_col)}{data_row}"

                    entry = {
                        "field_id": f"dual_{sheet_name}_row{data_row}",
                        "field_name": item_name,
                        "field_type": "dual_column_checkbox",
                        "pass_cell": {
                            "sheet": sheet_name,
                            "cell": pass_coord,
                            "row": data_row,
                            "column": pass_col,
                        },
                        "fail_cell": {
                            "sheet": sheet_name,
                            "cell": fail_coord,
                            "row": data_row,
                            "column": fail_col,
                        },
                        "remarks_cell": None,
                        "check_symbol": detected_symbol,
                    }

                    if remarks_col:
                        remarks_coord = f"{get_column_letter(remarks_col)}{data_row}"
                        entry["remarks_cell"] = {
                            "sheet": sheet_name,
                            "cell": remarks_coord,
                            "row": data_row,
                            "column": remarks_col,
                        }

                    dual_fields.append(entry)

            # Phase 3: 偵測已有的勾選符號
            detected = self._detect_check_symbol(ws, max_row, max_col)
            if detected:
                detected_symbol = detected

        return {
            "dual_column_fields": dual_fields,
            "check_symbol": detected_symbol,
            "total_items": len(dual_fields),
        }

    def _find_checkbox_header_rows(self, ws, max_row: int, max_col: int) -> list[dict]:
        """
        找到含有「合格/不合格」配對的表頭列

        支援的配對模式:
        - 合格 / 不合格
        - 正常 / 異常
        - 良好 / 不良
        - ○ / ×
        - Pass / Fail
        """
        PASS_KEYWORDS = ['合格', '正常', '良好', '良', '○', 'OK', 'Pass', 'PASS']
        FAIL_KEYWORDS = ['不合格', '異常', '不良', '×', 'NG', 'Fail', 'FAIL']
        REMARKS_KEYWORDS_HEADER = ['備註', '說明', '異常說明', '描述', '改善建議']
        ITEM_KEYWORDS_HEADER = ['檢查項目', '項目', '檢測項目', '量測項目', '項次']

        headers = []

        for row_idx in range(1, min(max_row, 20) + 1):  # 表頭通常在前 20 列
            pass_col = None
            fail_col = None
            remarks_col = None
            item_col = None

            for col_idx in range(1, max_col + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                val = str(cell.value).strip() if cell.value else ""

                if not val:
                    continue

                # 找合格欄
                if any(kw == val or kw in val for kw in PASS_KEYWORDS):
                    if not any(kw in val for kw in FAIL_KEYWORDS):  # 排除「不合格」
                        pass_col = col_idx

                # 找不合格欄
                if any(kw == val or kw in val for kw in FAIL_KEYWORDS):
                    fail_col = col_idx

                # 找備註欄
                if any(kw in val for kw in REMARKS_KEYWORDS_HEADER):
                    remarks_col = col_idx

                # 找項目欄
                if any(kw in val for kw in ITEM_KEYWORDS_HEADER):
                    item_col = col_idx

            # 必須同時有合格欄和不合格欄才算數
            if pass_col and fail_col:
                headers.append({
                    "row": row_idx,
                    "pass_col": pass_col,
                    "fail_col": fail_col,
                    "remarks_col": remarks_col,
                    "item_col": item_col or 1,  # 預設第 1 欄
                })

        return headers

    def _detect_check_symbol(self, ws, max_row: int, max_col: int) -> Optional[str]:
        """掃描表單中已有的勾選符號，學習該表的慣用符號"""
        KNOWN_SYMBOLS = ['✓', '✔', '○', '●', 'V', 'v', '√', '☑', '■']

        for row_idx in range(1, max_row + 1):
            for col_idx in range(1, max_col + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                val = str(cell.value).strip() if cell.value else ""
                if val in KNOWN_SYMBOLS:
                    return val

        return None

    async def _detect_checkbox_columns_word(self, content: bytes) -> dict:
        """偵測 Word 中的勾選雙欄結構"""
        doc = Document(io.BytesIO(content))
        dual_fields = []

        for table_idx, table in enumerate(doc.tables):
            # 找表頭列
            if len(table.rows) < 2:
                continue

            header_row = table.rows[0]
            cells_text = [cell.text.strip() for cell in header_row.cells]

            pass_col = None
            fail_col = None
            remarks_col = None
            item_col = None

            for ci, text in enumerate(cells_text):
                if text in ['合格', '正常', '良好', '○']:
                    pass_col = ci
                elif text in ['不合格', '異常', '不良', '×']:
                    fail_col = ci
                elif any(kw in text for kw in ['備註', '說明']):
                    remarks_col = ci
                elif any(kw in text for kw in ['檢查項目', '項目', '項次']):
                    item_col = ci

            if pass_col is None or fail_col is None:
                continue

            for ri in range(1, len(table.rows)):
                row = table.rows[ri]
                ic = item_col or 0
                item_name = row.cells[ic].text.strip() if ic < len(row.cells) else ""

                if not item_name:
                    continue

                entry = {
                    "field_id": f"dual_word_t{table_idx}_row{ri}",
                    "field_name": item_name,
                    "field_type": "dual_column_checkbox",
                    "pass_cell": {
                        "type": "table",
                        "table_index": table_idx,
                        "row_index": ri,
                        "cell_index": pass_col,
                    },
                    "fail_cell": {
                        "type": "table",
                        "table_index": table_idx,
                        "row_index": ri,
                        "cell_index": fail_col,
                    },
                    "remarks_cell": None,
                    "check_symbol": "✓",
                }

                if remarks_col is not None:
                    entry["remarks_cell"] = {
                        "type": "table",
                        "table_index": table_idx,
                        "row_index": ri,
                        "cell_index": remarks_col,
                    }

                dual_fields.append(entry)

        return {
            "dual_column_fields": dual_fields,
            "check_symbol": "✓",
            "total_items": len(dual_fields),
        }

    # ================================================================
    # 勾選回填引擎 (Sprint 2 Task 2.2 新增)
    # ================================================================

    async def auto_fill_with_checkboxes(
        self,
        file_content: bytes,
        file_name: str,
        field_map: list[dict],
        fill_values: list[dict],
        dual_column_fields: list[dict] = None,
        check_symbol: str = "✓",
    ) -> bytes:
        """
        增強版自動回填: 支援勾選式表單

        除了標準的文字/數值回填，額外處理 dual_column_checkbox 類型:
        - 合格 → 在 pass_cell 寫入勾選符號
        - 不合格 → 在 fail_cell 寫入勾選符號 + remarks_cell 寫入異常描述
        """
        file_type = file_name.split('.')[-1].lower()

        # 先做標準回填
        value_lookup = {fv["field_id"]: fv for fv in fill_values}
        field_lookup = {f["field_id"]: f for f in field_map}

        if file_type == 'xlsx':
            return await self._auto_fill_excel_enhanced(
                file_content, field_lookup, value_lookup,
                dual_column_fields or [], check_symbol
            )
        elif file_type == 'docx':
            return await self._auto_fill_word_enhanced(
                file_content, field_lookup, value_lookup,
                dual_column_fields or [], check_symbol
            )
        else:
            raise ValueError(f"不支援的檔案格式: {file_type}")

    async def _auto_fill_excel_enhanced(
        self,
        file_content: bytes,
        field_lookup: dict,
        value_lookup: dict,
        dual_column_fields: list[dict],
        check_symbol: str,
    ) -> bytes:
        """增強版 Excel 回填: 支援勾選"""
        wb = load_workbook(io.BytesIO(file_content))

        # Phase 1: 標準回填（文字/數值）
        for field_id, fv in value_lookup.items():
            field = field_lookup.get(field_id)
            if not field:
                continue

            # 跳過 dual_column_checkbox 類型（由 Phase 2 處理）
            if field.get("field_type") == "dual_column_checkbox":
                continue

            val_loc = field.get("value_location")
            if not val_loc:
                continue

            sheet_name = val_loc.get("sheet")
            cell_coord = val_loc.get("cell")
            if not sheet_name or not cell_coord or sheet_name not in wb.sheetnames:
                continue

            ws = wb[sheet_name]
            typed_value = self._convert_value(fv.get("value", fv) if isinstance(fv, dict) else fv, field.get("field_type", "text"))

            target_cell = ws[cell_coord]
            original_font = copy.copy(target_cell.font) if target_cell.font else None
            original_alignment = copy.copy(target_cell.alignment) if target_cell.alignment else None
            original_number_format = target_cell.number_format

            target_cell.value = typed_value

            if original_font:
                target_cell.font = original_font
            if original_alignment:
                target_cell.alignment = original_alignment
            if original_number_format:
                target_cell.number_format = original_number_format

        # Phase 2: 勾選回填
        for dual in dual_column_fields:
            field_id = dual.get("field_id", "")
            fv = value_lookup.get(field_id)
            if not fv:
                continue

            value = fv.get("value", fv) if isinstance(fv, dict) else fv
            is_pass = self._is_pass_value(str(value))

            pass_info = dual.get("pass_cell", {})
            fail_info = dual.get("fail_cell", {})
            remarks_info = dual.get("remarks_cell")
            symbol = dual.get("check_symbol", check_symbol)

            # 寫入勾選
            if is_pass:
                self._write_check_symbol(wb, pass_info, symbol)
                self._write_check_symbol(wb, fail_info, "")  # 清空另一欄
            else:
                self._write_check_symbol(wb, fail_info, symbol)
                self._write_check_symbol(wb, pass_info, "")

                # 不合格時寫入備註
                if remarks_info:
                    remarks_value = fv.get("remarks", "") if isinstance(fv, dict) else ""
                    if remarks_value:
                        self._write_check_symbol(wb, remarks_info, remarks_value)

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output.read()

    def _write_check_symbol(self, wb, cell_info: dict, value: str):
        """在指定儲存格寫入勾選符號，保留格式"""
        if not cell_info:
            return

        sheet_name = cell_info.get("sheet")
        cell_coord = cell_info.get("cell")

        if not sheet_name or not cell_coord or sheet_name not in wb.sheetnames:
            return

        ws = wb[sheet_name]
        target_cell = ws[cell_coord]

        original_font = copy.copy(target_cell.font) if target_cell.font else None
        original_alignment = copy.copy(target_cell.alignment) if target_cell.alignment else None

        target_cell.value = value

        if original_font:
            target_cell.font = original_font
        if original_alignment:
            target_cell.alignment = original_alignment

    def _is_pass_value(self, value: str) -> bool:
        """判斷值是否代表「合格/通過」"""
        v = value.strip().lower()
        return v in [
            'true', '1', '是', '合格', '正常', 'yes', 'ok', '通過',
            'pass', '良好', '良', '○',
        ]

    async def _auto_fill_word_enhanced(
        self,
        file_content: bytes,
        field_lookup: dict,
        value_lookup: dict,
        dual_column_fields: list[dict],
        check_symbol: str,
    ) -> bytes:
        """增強版 Word 回填: 支援勾選"""
        doc = Document(io.BytesIO(file_content))

        # Phase 1: 標準回填
        for field_id, fv in value_lookup.items():
            field = field_lookup.get(field_id)
            if not field:
                continue
            if field.get("field_type") == "dual_column_checkbox":
                continue

            val_loc = field.get("value_location")
            if not val_loc:
                continue

            value = fv.get("value", fv) if isinstance(fv, dict) else str(fv)
            loc_type = val_loc.get("type")

            if loc_type == "paragraph":
                para_idx = val_loc.get("paragraph_index")
                if para_idx is not None and para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    text = para.text
                    for sep in ['：', ':']:
                        if sep in text:
                            prefix = text.split(sep)[0] + sep
                            self._replace_paragraph_text_preserve_format(
                                para, f"{prefix} {value}"
                            )
                            break

            elif loc_type == "table":
                self._write_word_table_cell(doc, val_loc, str(value))

        # Phase 2: 勾選回填
        for dual in dual_column_fields:
            field_id = dual.get("field_id", "")
            fv = value_lookup.get(field_id)
            if not fv:
                continue

            value = fv.get("value", fv) if isinstance(fv, dict) else fv
            is_pass = self._is_pass_value(str(value))
            symbol = dual.get("check_symbol", check_symbol)

            if is_pass:
                self._write_word_table_cell(doc, dual.get("pass_cell", {}), symbol)
                self._write_word_table_cell(doc, dual.get("fail_cell", {}), "")
            else:
                self._write_word_table_cell(doc, dual.get("fail_cell", {}), symbol)
                self._write_word_table_cell(doc, dual.get("pass_cell", {}), "")

                if dual.get("remarks_cell"):
                    remarks_value = fv.get("remarks", "") if isinstance(fv, dict) else ""
                    if remarks_value:
                        self._write_word_table_cell(doc, dual["remarks_cell"], remarks_value)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _write_word_table_cell(self, doc, cell_info: dict, value: str):
        """在 Word 表格儲存格寫入值"""
        if not cell_info or cell_info.get("type") != "table":
            return

        table_idx = cell_info.get("table_index")
        row_idx = cell_info.get("row_index")
        cell_idx = cell_info.get("cell_index")

        if table_idx is None or row_idx is None or cell_idx is None:
            return
        if table_idx >= len(doc.tables):
            return

        table = doc.tables[table_idx]
        if row_idx >= len(table.rows):
            return

        row = table.rows[row_idx]
        if cell_idx >= len(row.cells):
            return

        cell = row.cells[cell_idx]
        if cell.paragraphs:
            self._replace_paragraph_text_preserve_format(cell.paragraphs[0], value)
        else:
            cell.text = value

    # ================================================================
    # 輔助方法
    # ================================================================

    def _replace_paragraph_text_preserve_format(self, paragraph, new_text: str):
        """替換段落文字但保留第一個 run 的格式"""
        if not paragraph.runs:
            paragraph.text = new_text
            return

        # 保留第一個 run 的格式
        first_run = paragraph.runs[0]

        # 清除所有 runs
        for run in paragraph.runs:
            run.text = ""

        # 在第一個 run 中設定新文字
        first_run.text = new_text

    def _convert_value(self, value, field_type: str):
        """根據欄位類型轉換值"""
        if value is None:
            return None

        if field_type == 'number':
            try:
                if '.' in str(value):
                    return float(value)
                return int(value)
            except (ValueError, TypeError):
                return str(value)
        elif field_type == 'checkbox':
            v = str(value).strip().lower()
            if v in ['true', '1', '是', '合格', '正常', 'yes', 'ok', '通過']:
                return '合格'
            elif v in ['false', '0', '否', '不合格', '異常', 'no', 'ng', '不通過']:
                return '不合格'
            return str(value)
        elif field_type == 'date':
            return str(value)
        else:
            return str(value)

    def _is_section_header(self, text: str) -> bool:
        """判斷文字是否為區段標題（而非可填入的欄位）"""
        text = text.strip()
        # 中文編號開頭的區段標題：一、二、三、... 或 （一）（二）...
        if re.match(r'^[一二三四五六七八九十]+[、．.]', text):
            return True
        if re.match(r'^[（(][一二三四五六七八九十]+[）)]', text):
            return True
        # 表格標題列的表頭欄位
        header_patterns = ['項次', '檢查項目', '檢查標準', '檢查要點', '量測項目',
                           '量測位置', '判定', '備註/異常說明', '備註']
        if text in header_patterns:
            return True
        return False

    def _is_non_field_item(self, text: str) -> bool:
        """判斷文字是否不應作為模板欄位（標題、表頭、注意事項等）"""
        text = text.strip()
        # 過長或過短的不太可能是欄位
        if len(text) > 30 or len(text) < 2:
            return True
        # 區段標題
        if self._is_section_header(text):
            return True
        # 常見非欄位文字
        non_field_patterns = [
            r'^注意事項',
            r'^簽核$',
            r'^\d+\.\s',  # 編號開頭的注意事項
            r'^□',         # 勾選項目描述
        ]
        return any(re.match(p, text) for p in non_field_patterns)
