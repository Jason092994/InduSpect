"""
Sprint 2 測試: 勾選式表單偵測 + 回填

測試範圍:
- Task 2.1: detect_checkbox_columns() 偵測合格/不合格雙欄結構
- Task 2.2: auto_fill_with_checkboxes() 勾選符號回填
"""

import sys
import os
import asyncio
import io
from datetime import datetime

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
os.environ.setdefault("GEMINI_API_KEY", "test-key")

from openpyxl import Workbook
from docx import Document

from app.services.form_fill import FormFillService


class TestResults:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.errors = []

    def check(self, condition, name, detail=""):
        if condition:
            self.passed += 1
            print(f"  ✅ {name}")
        else:
            self.failed += 1
            self.errors.append(f"{name}: {detail}")
            print(f"  ❌ {name} — {detail}")


def create_test_excel_with_checkboxes() -> bytes:
    """建立含有「合格/不合格」雙欄的測試 Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    # 表頭
    ws["A1"] = "電氣設備定期檢查表"
    ws["A2"] = "設備名稱"
    ws["B2"] = ""  # value cell
    ws["C2"] = "檢查日期"
    ws["D2"] = ""

    # 勾選表頭列 (Row 4)
    ws["A4"] = "檢查項目"
    ws["B4"] = "量測值"
    ws["C4"] = "合格"
    ws["D4"] = "不合格"
    ws["E4"] = "備註"

    # 資料列
    ws["A5"] = "絕緣電阻 R相"
    ws["B5"] = ""
    ws["C5"] = ""  # 合格打勾處
    ws["D5"] = ""  # 不合格打勾處
    ws["E5"] = ""

    ws["A6"] = "絕緣電阻 S相"
    ws["B6"] = ""
    ws["C6"] = ""
    ws["D6"] = ""
    ws["E6"] = ""

    ws["A7"] = "接地電阻"
    ws["B7"] = ""
    ws["C7"] = ""
    ws["D7"] = ""
    ws["E7"] = ""

    ws["A8"] = "漏電斷路器動作時間"
    ws["B8"] = ""
    ws["C8"] = ""
    ws["D8"] = ""
    ws["E8"] = ""

    # 已有勾選符號的欄位（用於符號偵測）
    ws["C5"] = "✓"  # 模擬已經有人打勾

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.read()


def create_test_excel_with_circle_symbols() -> bytes:
    """建立使用 ○/× 符號的測試 Excel"""
    wb = Workbook()
    ws = wb.active

    ws["A1"] = "項目"
    ws["B1"] = "正常"
    ws["C1"] = "異常"
    ws["D1"] = "說明"

    ws["A2"] = "外觀檢查"
    ws["B2"] = "○"
    ws["C2"] = ""
    ws["D2"] = ""

    ws["A3"] = "操作功能"
    ws["B3"] = ""
    ws["C3"] = ""
    ws["D3"] = ""

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.read()


def create_test_word_with_checkboxes() -> bytes:
    """建立含有勾選表格的測試 Word"""
    doc = Document()
    doc.add_heading("消防設備定期檢查表", level=1)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # 表頭
    table.rows[0].cells[0].text = "檢查項目"
    table.rows[0].cells[1].text = "合格"
    table.rows[0].cells[2].text = "不合格"
    table.rows[0].cells[3].text = "備註"

    # 資料
    table.rows[1].cells[0].text = "滅火器壓力"
    table.rows[2].cells[0].text = "緊急照明"
    table.rows[3].cells[0].text = "偵煙探測器"

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


async def test_detect_checkbox_excel():
    """測試 Excel 勾選偵測"""
    print("\n" + "="*60)
    print("TEST: Task 2.1 — detect_checkbox_columns (Excel)")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    content = create_test_excel_with_checkboxes()
    result = await service.detect_checkbox_columns(content, "test.xlsx")

    dual_fields = result["dual_column_fields"]
    symbol = result["check_symbol"]

    # Test 1: 偵測到雙欄結構
    results.check(
        len(dual_fields) >= 3,
        f"偵測到至少 3 個勾選項目",
        f"實際: {len(dual_fields)}"
    )

    # Test 2: 偵測到已有的勾選符號
    results.check(
        symbol == "✓",
        f"偵測到勾選符號 = ✓",
        f"實際: {symbol}"
    )

    # Test 3: 每個項目都有 pass_cell 和 fail_cell
    for df in dual_fields:
        results.check(
            df.get("pass_cell") is not None,
            f"「{df['field_name']}」有 pass_cell"
        )
        results.check(
            df.get("fail_cell") is not None,
            f"「{df['field_name']}」有 fail_cell"
        )

    # Test 4: 備註欄位正確偵測
    has_remarks = any(df.get("remarks_cell") is not None for df in dual_fields)
    results.check(
        has_remarks,
        "至少一個項目有 remarks_cell"
    )

    # Test 5: field_type 正確
    for df in dual_fields:
        results.check(
            df["field_type"] == "dual_column_checkbox",
            f"「{df['field_name']}」field_type = dual_column_checkbox",
            f"實際: {df['field_type']}"
        )
        break  # 只測第一個

    # Test 6: 項目名稱正確
    names = [df["field_name"] for df in dual_fields]
    results.check(
        "絕緣電阻 R相" in names,
        "包含「絕緣電阻 R相」",
        f"所有名稱: {names}"
    )

    return results


async def test_detect_checkbox_circle():
    """測試 ○/× 符號的偵測"""
    print("\n" + "="*60)
    print("TEST: Task 2.1 — detect_checkbox (○/× 符號)")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    content = create_test_excel_with_circle_symbols()
    result = await service.detect_checkbox_columns(content, "test_circle.xlsx")

    dual_fields = result["dual_column_fields"]
    symbol = result["check_symbol"]

    results.check(
        len(dual_fields) >= 1,
        f"偵測到至少 1 個勾選項目（○/× 模式）",
        f"實際: {len(dual_fields)}"
    )

    results.check(
        symbol == "○",
        f"偵測到勾選符號 = ○",
        f"實際: {symbol}"
    )

    return results


async def test_detect_checkbox_word():
    """測試 Word 勾選偵測"""
    print("\n" + "="*60)
    print("TEST: Task 2.1 — detect_checkbox_columns (Word)")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    content = create_test_word_with_checkboxes()
    result = await service.detect_checkbox_columns(content, "test.docx")

    dual_fields = result["dual_column_fields"]

    results.check(
        len(dual_fields) == 3,
        f"偵測到 3 個勾選項目",
        f"實際: {len(dual_fields)}"
    )

    names = [df["field_name"] for df in dual_fields]
    results.check(
        "滅火器壓力" in names,
        "包含「滅火器壓力」",
        f"名稱: {names}"
    )
    results.check(
        "緊急照明" in names,
        "包含「緊急照明」"
    )

    return results


async def test_auto_fill_checkboxes_excel():
    """測試 Excel 勾選回填"""
    print("\n" + "="*60)
    print("TEST: Task 2.2 — auto_fill_with_checkboxes (Excel)")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    content = create_test_excel_with_checkboxes()

    # 先偵測雙欄結構
    detect_result = await service.detect_checkbox_columns(content, "test.xlsx")
    dual_fields = detect_result["dual_column_fields"]
    check_symbol = detect_result["check_symbol"]

    # 準備 fill_values
    fill_values = []
    for df in dual_fields:
        if "R相" in df["field_name"]:
            fill_values.append({"field_id": df["field_id"], "value": "合格"})
        elif "S相" in df["field_name"]:
            fill_values.append({"field_id": df["field_id"], "value": "合格"})
        elif "接地" in df["field_name"]:
            fill_values.append({
                "field_id": df["field_id"],
                "value": "不合格",
                "remarks": "接地電阻 85.2Ω，偏高"
            })
        elif "漏電" in df["field_name"]:
            fill_values.append({"field_id": df["field_id"], "value": "合格"})

    # 執行回填
    filled_bytes = await service.auto_fill_with_checkboxes(
        file_content=content,
        file_name="test.xlsx",
        field_map=[],  # 不做標準回填，只測勾選
        fill_values=fill_values,
        dual_column_fields=dual_fields,
        check_symbol=check_symbol,
    )

    results.check(
        filled_bytes is not None and len(filled_bytes) > 0,
        "回填後產生了有效的 Excel 檔案"
    )

    # 驗證回填結果
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(filled_bytes))
    ws = wb.active

    # R相應該在合格欄打勾
    r_pass = ws["C5"].value
    r_fail = ws["D5"].value
    results.check(
        r_pass == check_symbol,
        f"R相合格欄 = {check_symbol}",
        f"實際: {r_pass}"
    )
    results.check(
        r_fail == "" or r_fail is None,
        "R相不合格欄為空",
        f"實際: {r_fail}"
    )

    # 接地應該在不合格欄打勾
    ground_pass = ws["C7"].value
    ground_fail = ws["D7"].value
    results.check(
        ground_fail == check_symbol,
        f"接地電阻不合格欄 = {check_symbol}",
        f"實際: {ground_fail}"
    )
    results.check(
        ground_pass == "" or ground_pass is None,
        "接地電阻合格欄為空",
        f"實際: {ground_pass}"
    )

    # 接地備註應有異常說明
    ground_remarks = ws["E7"].value
    results.check(
        ground_remarks is not None and "85.2" in str(ground_remarks),
        "接地電阻備註包含異常數值",
        f"實際: {ground_remarks}"
    )

    return results


async def test_auto_fill_checkboxes_word():
    """測試 Word 勾選回填"""
    print("\n" + "="*60)
    print("TEST: Task 2.2 — auto_fill_with_checkboxes (Word)")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    content = create_test_word_with_checkboxes()
    detect_result = await service.detect_checkbox_columns(content, "test.docx")
    dual_fields = detect_result["dual_column_fields"]

    fill_values = []
    for df in dual_fields:
        if "滅火器" in df["field_name"]:
            fill_values.append({"field_id": df["field_id"], "value": "合格"})
        elif "緊急照明" in df["field_name"]:
            fill_values.append({"field_id": df["field_id"], "value": "不合格", "remarks": "電池失效"})
        elif "偵煙" in df["field_name"]:
            fill_values.append({"field_id": df["field_id"], "value": "合格"})

    filled_bytes = await service.auto_fill_with_checkboxes(
        file_content=content,
        file_name="test.docx",
        field_map=[],
        fill_values=fill_values,
        dual_column_fields=dual_fields,
        check_symbol="✓",
    )

    results.check(
        filled_bytes is not None and len(filled_bytes) > 0,
        "回填後產生了有效的 Word 檔案"
    )

    # 驗證回填結果
    doc = Document(io.BytesIO(filled_bytes))
    table = doc.tables[0]

    # 滅火器合格欄
    fire_pass = table.rows[1].cells[1].text.strip()
    fire_fail = table.rows[1].cells[2].text.strip()
    results.check(
        fire_pass == "✓",
        f"滅火器合格欄 = ✓",
        f"實際: '{fire_pass}'"
    )
    results.check(
        fire_fail == "",
        "滅火器不合格欄為空",
        f"實際: '{fire_fail}'"
    )

    # 緊急照明不合格欄
    light_pass = table.rows[2].cells[1].text.strip()
    light_fail = table.rows[2].cells[2].text.strip()
    results.check(
        light_fail == "✓",
        f"緊急照明不合格欄 = ✓",
        f"實際: '{light_fail}'"
    )
    results.check(
        light_pass == "",
        "緊急照明合格欄為空",
        f"實際: '{light_pass}'"
    )

    return results


async def main():
    print("="*60)
    print(f"Sprint 2 測試報告 — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("="*60)

    all_results = []
    all_results.append(await test_detect_checkbox_excel())
    all_results.append(await test_detect_checkbox_circle())
    all_results.append(await test_detect_checkbox_word())
    all_results.append(await test_auto_fill_checkboxes_excel())
    all_results.append(await test_auto_fill_checkboxes_word())

    total_passed = sum(r.passed for r in all_results)
    total_failed = sum(r.failed for r in all_results)
    total = total_passed + total_failed

    print("\n" + "="*60)
    print(f"Sprint 2 總結")
    print("="*60)
    print(f"總測試數: {total}")
    print(f"通過: {total_passed} ✅")
    print(f"失敗: {total_failed} ❌")
    print(f"通過率: {total_passed/total*100:.1f}%" if total > 0 else "N/A")

    if total_failed > 0:
        print("\n失敗清單:")
        for r in all_results:
            for e in r.errors:
                print(f"  ❌ {e}")

    return total_failed == 0


if __name__ == "__main__":
    success = asyncio.run(main())
    sys.exit(0 if success else 1)
