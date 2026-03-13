"""
Sprint 2 測試: Task 2.3 — 照片自動插入報告

測試範圍:
- insert_photos_into_report() Excel 照片插入
- insert_photos_into_report() Word 照片插入
- _prepare_photo_for_insert() 照片處理
"""

import sys
import os
import asyncio
import io
import base64
from datetime import datetime

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
os.environ.setdefault("GEMINI_API_KEY", "test-key")

from openpyxl import Workbook, load_workbook
from docx import Document
from PIL import Image as PILImage

from app.services.form_fill import FormFillService


class TestResults:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.errors = []

    def check(self, condition, name, detail=""):
        if condition:
            self.passed += 1
            print(f"  ✅ {name}")
        else:
            self.failed += 1
            self.errors.append(f"{name}: {detail}")
            print(f"  ❌ {name} — {detail}")


def create_test_photo(width=800, height=600, color=(100, 150, 200)) -> bytes:
    """建立測試用的假照片 (JPEG bytes)"""
    img = PILImage.new('RGB', (width, height), color=color)
    output = io.BytesIO()
    img.save(output, format='JPEG', quality=80)
    output.seek(0)
    return output.read()


def create_test_photo_base64(width=800, height=600, color=(100, 150, 200)) -> str:
    """建立測試用的假照片 (base64 字串)"""
    photo_bytes = create_test_photo(width, height, color)
    return base64.b64encode(photo_bytes).decode('utf-8')


def create_test_excel() -> bytes:
    """建立簡單的測試 Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "電氣設備定期檢查表"
    ws["A2"] = "設備名稱"
    ws["B2"] = "B棟1F配電盤"
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.read()


def create_test_word() -> bytes:
    """建立簡單的測試 Word"""
    doc = Document()
    doc.add_heading("消防設備定期檢查表", level=1)
    doc.add_paragraph("設備名稱: B棟消防栓")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


async def test_prepare_photo():
    """測試照片前處理"""
    print("\n" + "="*60)
    print("TEST: Task 2.3 — _prepare_photo_for_insert()")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    # Test 1: base64 照片處理
    b64 = create_test_photo_base64(1200, 900)
    binding = {"task_id": "t1", "photo_base64": b64}
    photo_io = service._prepare_photo_for_insert(binding)
    results.check(
        photo_io is not None,
        "base64 照片成功處理"
    )

    if photo_io:
        img = PILImage.open(photo_io)
        results.check(
            img.size[0] <= 600,
            f"照片寬度 <= 600px",
            f"實際: {img.size[0]}"
        )
        results.check(
            img.size[1] <= 450,
            f"照片高度 <= 450px",
            f"實際: {img.size[1]}"
        )

    # Test 2: bytes 照片處理
    photo_bytes = create_test_photo(400, 300)
    binding2 = {"task_id": "t2", "photo_bytes": photo_bytes}
    photo_io2 = service._prepare_photo_for_insert(binding2)
    results.check(
        photo_io2 is not None,
        "bytes 照片成功處理"
    )

    # Test 3: 無照片資料
    binding3 = {"task_id": "t3"}
    photo_io3 = service._prepare_photo_for_insert(binding3)
    results.check(
        photo_io3 is None,
        "無照片資料回傳 None"
    )

    # Test 4: base64 帶 data URI 前綴
    b64_with_prefix = f"data:image/jpeg;base64,{create_test_photo_base64(200, 150)}"
    binding4 = {"task_id": "t4", "photo_base64": b64_with_prefix}
    photo_io4 = service._prepare_photo_for_insert(binding4)
    results.check(
        photo_io4 is not None,
        "帶 data URI 前綴的 base64 成功處理"
    )

    # Test 5: RGBA 照片（例如 PNG with transparency）
    rgba_img = PILImage.new('RGBA', (300, 200), color=(100, 150, 200, 128))
    rgba_buf = io.BytesIO()
    rgba_img.save(rgba_buf, format='PNG')
    rgba_bytes = rgba_buf.getvalue()
    binding5 = {"task_id": "t5", "photo_bytes": rgba_bytes}
    photo_io5 = service._prepare_photo_for_insert(binding5)
    results.check(
        photo_io5 is not None,
        "RGBA 照片轉換為 RGB 成功"
    )

    # Test 6: 照片大小合理 (< 500KB)
    large_b64 = create_test_photo_base64(2000, 1500)
    binding6 = {"task_id": "t6", "photo_base64": large_b64}
    photo_io6 = service._prepare_photo_for_insert(binding6)
    if photo_io6:
        size_kb = len(photo_io6.read()) / 1024
        photo_io6.seek(0)
        results.check(
            size_kb <= 500,
            f"大照片壓縮後 <= 500KB",
            f"實際: {size_kb:.1f}KB"
        )

    return results


async def test_insert_photos_excel():
    """測試 Excel 照片插入"""
    print("\n" + "="*60)
    print("TEST: Task 2.3 — insert_photos_into_report (Excel)")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    excel_content = create_test_excel()

    photo_bindings = [
        {
            "task_id": "task_1",
            "display_name": "絕緣電阻測量",
            "photo_base64": create_test_photo_base64(800, 600, (200, 100, 100)),
            "capture_time": "2026-03-13 14:30",
            "sequence": 1,
        },
        {
            "task_id": "task_2",
            "display_name": "接地電阻測量",
            "photo_base64": create_test_photo_base64(800, 600, (100, 200, 100)),
            "capture_time": "2026-03-13 14:45",
            "sequence": 2,
        },
        {
            "task_id": "task_3",
            "display_name": "漏電斷路器動作測試",
            "photo_base64": create_test_photo_base64(800, 600, (100, 100, 200)),
            "capture_time": "2026-03-13 15:00",
            "sequence": 3,
        },
    ]

    result_bytes = await service.insert_photos_into_report(
        file_content=excel_content,
        file_name="test.xlsx",
        photo_bindings=photo_bindings,
    )

    # Test 1: 產生有效檔案
    results.check(
        result_bytes is not None and len(result_bytes) > 0,
        "產生有效的 Excel 檔案"
    )

    # Test 2: 驗證工作表
    wb = load_workbook(io.BytesIO(result_bytes))
    results.check(
        "照片附件" in wb.sheetnames,
        "新增了「照片附件」工作表",
        f"工作表: {wb.sheetnames}"
    )

    ws = wb["照片附件"]

    # Test 3: 標題
    results.check(
        ws["A1"].value == "現場照片記錄",
        "標題 = 「現場照片記錄」",
        f"實際: {ws['A1'].value}"
    )

    # Test 4: 表頭
    results.check(
        ws.cell(row=2, column=1).value == "編號",
        "表頭第1欄 = 編號"
    )
    results.check(
        ws.cell(row=2, column=2).value == "檢查項目",
        "表頭第2欄 = 檢查項目"
    )
    results.check(
        ws.cell(row=2, column=3).value == "現場照片",
        "表頭第3欄 = 現場照片"
    )

    # Test 5: 資料列內容
    results.check(
        ws.cell(row=3, column=2).value == "絕緣電阻測量",
        "第1筆檢查項目 = 絕緣電阻測量",
        f"實際: {ws.cell(row=3, column=2).value}"
    )
    results.check(
        ws.cell(row=4, column=2).value == "接地電阻測量",
        "第2筆檢查項目 = 接地電阻測量"
    )
    results.check(
        ws.cell(row=5, column=2).value == "漏電斷路器動作測試",
        "第3筆檢查項目 = 漏電斷路器動作測試"
    )

    # Test 6: 拍攝時間
    results.check(
        ws.cell(row=3, column=4).value == "2026-03-13 14:30",
        "第1筆拍攝時間正確"
    )

    # Test 7: 有照片（檢查圖片數量）
    image_count = len(ws._images)
    results.check(
        image_count == 3,
        f"插入了 3 張照片",
        f"實際: {image_count}"
    )

    # Test 8: 原始工作表未被影響
    results.check(
        "Sheet1" in wb.sheetnames,
        "原始 Sheet1 保留"
    )
    results.check(
        wb["Sheet1"]["A1"].value == "電氣設備定期檢查表",
        "原始資料未被修改"
    )

    return results


async def test_insert_photos_word():
    """測試 Word 照片插入"""
    print("\n" + "="*60)
    print("TEST: Task 2.3 — insert_photos_into_report (Word)")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    word_content = create_test_word()

    photo_bindings = [
        {
            "task_id": "task_1",
            "display_name": "滅火器壓力檢查",
            "photo_base64": create_test_photo_base64(800, 600, (200, 150, 100)),
            "capture_time": "2026-03-13 10:00",
            "sequence": 1,
        },
        {
            "task_id": "task_2",
            "display_name": "緊急照明功能測試",
            "photo_base64": create_test_photo_base64(800, 600, (100, 150, 200)),
            "capture_time": "2026-03-13 10:15",
            "sequence": 2,
        },
    ]

    result_bytes = await service.insert_photos_into_report(
        file_content=word_content,
        file_name="test.docx",
        photo_bindings=photo_bindings,
    )

    # Test 1: 產生有效檔案
    results.check(
        result_bytes is not None and len(result_bytes) > 0,
        "產生有效的 Word 檔案"
    )

    # Test 2: 驗證內容
    doc = Document(io.BytesIO(result_bytes))

    # 原始內容保留
    results.check(
        doc.paragraphs[0].text == "消防設備定期檢查表",
        "原始標題保留",
        f"實際: {doc.paragraphs[0].text}"
    )

    # Test 3: 照片記錄標題
    all_text = "\n".join([p.text for p in doc.paragraphs])
    results.check(
        "照片記錄" in all_text,
        "包含「照片記錄」章節標題"
    )

    # Test 4: 照片項目名稱
    results.check(
        "滅火器壓力檢查" in all_text,
        "包含「滅火器壓力檢查」項目名稱"
    )
    results.check(
        "緊急照明功能測試" in all_text,
        "包含「緊急照明功能測試」項目名稱"
    )

    # Test 5: 拍攝時間
    results.check(
        "2026-03-13 10:00" in all_text,
        "包含拍攝時間"
    )

    # Test 6: 確認有嵌入圖片（檢查 inline shapes）
    # Word 中的圖片是 inline shapes
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    results.check(
        image_count >= 2,
        f"插入了至少 2 張照片",
        f"實際: {image_count}"
    )

    # Test 7: 檔案大小合理（原始 + 2張壓縮照片不應超過 2MB）
    size_kb = len(result_bytes) / 1024
    results.check(
        size_kb < 2048,
        f"檔案大小合理 (< 2MB)",
        f"實際: {size_kb:.1f}KB"
    )

    return results


async def test_insert_photos_empty():
    """測試邊界情況"""
    print("\n" + "="*60)
    print("TEST: Task 2.3 — 邊界情況")
    print("="*60)

    results = TestResults()
    service = FormFillService()

    # Test 1: 無效照片資料（應跳過，不崩潰）
    excel_content = create_test_excel()
    photo_bindings = [
        {
            "task_id": "task_bad",
            "display_name": "無效照片",
            "photo_base64": "not-valid-base64!!!",
            "sequence": 1,
        },
        {
            "task_id": "task_good",
            "display_name": "有效照片",
            "photo_base64": create_test_photo_base64(400, 300),
            "capture_time": "2026-03-13 11:00",
            "sequence": 2,
        },
    ]

    result_bytes = await service.insert_photos_into_report(
        file_content=excel_content,
        file_name="test.xlsx",
        photo_bindings=photo_bindings,
    )

    results.check(
        result_bytes is not None and len(result_bytes) > 0,
        "無效照片不影響整體處理"
    )

    wb = load_workbook(io.BytesIO(result_bytes))
    ws = wb["照片附件"]
    # 只有有效的那張照片被插入
    image_count = len(ws._images)
    results.check(
        image_count == 1,
        f"只插入有效照片（跳過無效的）",
        f"實際圖片數: {image_count}"
    )

    # Test 2: 不支援的檔案類型
    try:
        await service.insert_photos_into_report(
            file_content=b"test",
            file_name="test.pdf",
            photo_bindings=[],
        )
        results.check(False, "不支援的檔案類型應拋出 ValueError")
    except ValueError:
        results.check(True, "不支援的檔案類型正確拋出 ValueError")

    # Test 3: 排序功能（sequence 亂序）
    excel_content2 = create_test_excel()
    bindings_unsorted = [
        {
            "task_id": "t3",
            "display_name": "第三項",
            "photo_base64": create_test_photo_base64(200, 150),
            "sequence": 3,
        },
        {
            "task_id": "t1",
            "display_name": "第一項",
            "photo_base64": create_test_photo_base64(200, 150),
            "sequence": 1,
        },
        {
            "task_id": "t2",
            "display_name": "第二項",
            "photo_base64": create_test_photo_base64(200, 150),
            "sequence": 2,
        },
    ]
    result_sorted = await service.insert_photos_into_report(
        file_content=excel_content2,
        file_name="test.xlsx",
        photo_bindings=bindings_unsorted,
    )
    wb2 = load_workbook(io.BytesIO(result_sorted))
    ws2 = wb2["照片附件"]
    results.check(
        ws2.cell(row=3, column=2).value == "第一項",
        "照片按 sequence 排序（第一項在最前）",
        f"實際: {ws2.cell(row=3, column=2).value}"
    )
    results.check(
        ws2.cell(row=4, column=2).value == "第二項",
        "第二項排在第二"
    )
    results.check(
        ws2.cell(row=5, column=2).value == "第三項",
        "第三項排在最後"
    )

    return results


async def main():
    print("="*60)
    print(f"Sprint 2 Task 2.3 測試報告 — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("="*60)

    all_results = []
    all_results.append(await test_prepare_photo())
    all_results.append(await test_insert_photos_excel())
    all_results.append(await test_insert_photos_word())
    all_results.append(await test_insert_photos_empty())

    total_passed = sum(r.passed for r in all_results)
    total_failed = sum(r.failed for r in all_results)
    total = total_passed + total_failed

    print("\n" + "="*60)
    print(f"Task 2.3 總結")
    print("="*60)
    print(f"總測試數: {total}")
    print(f"通過: {total_passed} ✅")
    print(f"失敗: {total_failed} ❌")
    print(f"通過率: {total_passed/total*100:.1f}%" if total > 0 else "N/A")

    if total_failed > 0:
        print("\n失敗清單:")
        for r in all_results:
            for e in r.errors:
                print(f"  ❌ {e}")

    return total_failed == 0


if __name__ == "__main__":
    success = asyncio.run(main())
    sys.exit(0 if success else 1)
