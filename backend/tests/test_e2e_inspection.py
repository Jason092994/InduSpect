"""
Sprint 6 Task 6.1: E2E 自動化測試

6 個端到端測試案例:
1. 電氣設備定檢 — 全合格場景
2. 電氣設備定檢 — 有異常場景
3. 消防設備定檢 — Word 格式
4. 勾選式表單
5. 照片插入
6. 歷史資料帶入
"""

import sys
import os
import asyncio
import tempfile
import io

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
os.environ.setdefault("GEMINI_API_KEY", "test-key")

from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
from PIL import Image as PILImage

from app.services.form_fill import FormFillService
from app.services.history_service import HistoryService


class TestResults:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.errors = []

    def check(self, condition, name, detail=""):
        if condition:
            self.passed += 1
            print(f"  PASS {name}")
        else:
            self.failed += 1
            self.errors.append(f"{name}: {detail}")
            print(f"  FAIL {name} -- {detail}")


# ============================================================
# 工具函式: 建立測試用 Excel/Word 表單
# ============================================================

def create_electrical_excel_form() -> bytes:
    """建立電氣設備定檢 Excel 表單"""
    wb = Workbook()
    ws = wb.active
    ws.title = "電氣設備定檢"

    # 基本資訊區 (uses ':' keyword for label detection)
    ws['A1'] = '設備名稱：'
    ws['B1'] = ''
    ws['A2'] = '設備編號：'
    ws['B2'] = ''
    ws['A3'] = '檢查日期：'
    ws['B3'] = ''
    ws['A4'] = '檢查人員：'
    ws['B4'] = ''

    # 檢測項目 (uses keywords: 溫度/壓力/電壓/電流/讀數/測量 etc.)
    ws['A6'] = '絕緣電阻R相 讀數：'
    ws['B6'] = ''
    ws['C6'] = '判定結果'
    ws['D6'] = '備註'

    ws['A7'] = '絕緣電阻S相 讀數：'
    ws['B7'] = ''
    ws['C7'] = '判定結果'
    ws['D7'] = '備註'

    ws['A8'] = '接地電阻 測量：'
    ws['B8'] = ''
    ws['C8'] = '判定結果'
    ws['D8'] = '備註'

    ws['A9'] = '漏電斷路器 測量：'
    ws['B9'] = ''
    ws['C9'] = '判定結果'
    ws['D9'] = '備註'

    # 結論區
    ws['A11'] = '綜合結論：'
    ws['B11'] = ''

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def create_fire_word_form() -> bytes:
    """建立消防設備定檢 Word 表單"""
    doc = Document()
    doc.add_heading('消防安全設備定期檢查表', level=1)

    # 基本資訊段落
    doc.add_paragraph('設備名稱：')
    doc.add_paragraph('檢查日期：')
    doc.add_paragraph('檢查人員：')

    # 檢查項目表格
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # 表頭
    header_cells = table.rows[0].cells
    header_cells[0].text = '檢查項目'
    header_cells[1].text = '量測值'
    header_cells[2].text = '結果'
    header_cells[3].text = '備註'

    # 資料行
    items = [
        '滅火器壓力',
        '緊急照明持續時間',
        '偵煙探測器靈敏度',
        '消防栓放水壓力',
    ]
    for i, item in enumerate(items, 1):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = ''
        table.rows[i].cells[2].text = ''
        table.rows[i].cells[3].text = ''

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def create_checkbox_excel_form() -> bytes:
    """建立勾選式 Excel 表單 (合格/不合格 雙欄)"""
    wb = Workbook()
    ws = wb.active
    ws.title = "勾選式檢查表"

    # 表頭
    ws['A1'] = '檢查項目'
    ws['B1'] = '合格'
    ws['C1'] = '不合格'
    ws['D1'] = '備註'

    # 檢查項目
    items = ['外觀檢查', '絕緣測試', '接地測試', '功能測試', '安全裝置']
    for i, item in enumerate(items, 2):
        ws[f'A{i}'] = item
        ws[f'B{i}'] = ''
        ws[f'C{i}'] = ''
        ws[f'D{i}'] = ''

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def create_test_photo(width=100, height=100, color='red') -> bytes:
    """建立測試用照片 (PIL)"""
    img = PILImage.new('RGB', (width, height), color=color)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf.read()


# ============================================================
# TEST 1: 電氣設備定檢 — 全合格場景
# ============================================================

async def test_case_1_electrical_all_pass():
    """電氣設備定檢 — 全合格場景"""
    print("\n" + "=" * 60)
    print("TEST 1: 電氣設備定檢 — 全合格場景")
    print("=" * 60)

    results = TestResults()
    service = FormFillService()

    # Step 1: 建立測試表單
    file_content = create_electrical_excel_form()
    file_name = "電氣設備定檢表.xlsx"

    # Step 2: analyze_structure
    structure = await service.analyze_structure(file_content, file_name)
    results.check(structure["success"], "analyze_structure 成功")
    results.check(structure["total_fields"] > 0, "偵測到欄位",
                  f"total_fields={structure['total_fields']}")

    field_map = structure["field_map"]

    # Step 3: generate_photo_tasks
    tasks_result = await service.generate_photo_tasks(field_map)
    photo_tasks = tasks_result["photo_tasks"]
    results.check(len(photo_tasks) > 0, "產生拍照任務",
                  f"tasks={len(photo_tasks)}")
    results.check(len(tasks_result["basic_info_fields"]) > 0, "偵測到基本資訊欄位",
                  f"basic={len(tasks_result['basic_info_fields'])}")

    # Step 4: 模擬 AI 結果 (全合格)
    simulated_ai_results = [
        {
            "task_id": photo_tasks[0]["task_id"] if len(photo_tasks) > 0 else "photo_001",
            "ai_result": {
                "readings": [{"label": "絕緣電阻", "value": 52.3, "unit": "MΩ"}],
                "is_anomaly": False,
                "condition_assessment": "合格",
                "anomaly_description": "",
                "summary": "絕緣電阻正常",
            },
            "value_field_ids": photo_tasks[0]["value_field_ids"] if len(photo_tasks) > 0 else [],
            "judgment_field_ids": photo_tasks[0]["judgment_field_ids"] if len(photo_tasks) > 0 else [],
            "remarks_field_ids": photo_tasks[0]["remarks_field_ids"] if len(photo_tasks) > 0 else [],
        },
    ]

    # Add more bindings for other tasks
    for i in range(1, min(len(photo_tasks), 4)):
        task = photo_tasks[i]
        simulated_ai_results.append({
            "task_id": task["task_id"],
            "ai_result": {
                "readings": [{"label": task["display_name"], "value": 45.0, "unit": "Ω"}],
                "is_anomaly": False,
                "condition_assessment": "合格",
                "anomaly_description": "",
                "summary": "正常",
            },
            "value_field_ids": task["value_field_ids"],
            "judgment_field_ids": task["judgment_field_ids"],
            "remarks_field_ids": task["remarks_field_ids"],
        })

    # Step 5: precision_map_fields
    general_info = [{
        "equipment_name": "A棟配電盤",
        "equipment_id": "EQ-001",
        "inspection_date": "2026-03-13",
        "inspector_name": "王小明",
        "location": "A棟1F",
    }]

    mapping_result = await service.precision_map_fields(
        field_map=field_map,
        inspection_results=general_info,
        photo_task_bindings=simulated_ai_results,
    )
    results.check(mapping_result["success"], "precision_map_fields 成功")

    mappings = mapping_result["mappings"]
    results.check(len(mappings) > 0, "產生映射結果",
                  f"mappings={len(mappings)}")

    # Step 6: auto_fill
    fill_values = [{"field_id": m["field_id"], "value": m["suggested_value"]}
                   for m in mappings]

    filled_bytes = await service.auto_fill(file_content, file_name, field_map, fill_values)
    results.check(len(filled_bytes) > 0, "auto_fill 產生非空檔案")

    # 驗證回填結果
    wb = load_workbook(io.BytesIO(filled_bytes))
    ws = wb.active
    filled_count = 0
    for row in range(1, 13):
        for col in range(1, 5):
            val = ws.cell(row=row, column=col).value
            if val is not None and str(val).strip():
                filled_count += 1

    results.check(filled_count > 4, "回填後多個儲存格有值",
                  f"filled_count={filled_count}")

    # 驗證合格判定 — 檢查映射中是否有合格
    pass_mappings = [m for m in mappings if '合格' in str(m.get("suggested_value", ""))]
    results.check(len(pass_mappings) > 0, "全部判定為合格",
                  f"pass_mappings={len(pass_mappings)}")

    return results


# ============================================================
# TEST 2: 電氣設備定檢 — 有異常場景
# ============================================================

async def test_case_2_electrical_with_anomaly():
    """電氣設備定檢 — 有異常場景"""
    print("\n" + "=" * 60)
    print("TEST 2: 電氣設備定檢 — 有異常場景")
    print("=" * 60)

    results = TestResults()
    service = FormFillService()

    file_content = create_electrical_excel_form()
    file_name = "電氣設備定檢表.xlsx"

    # analyze_structure
    structure = await service.analyze_structure(file_content, file_name)
    field_map = structure["field_map"]
    results.check(structure["success"], "analyze_structure 成功")

    # generate_photo_tasks
    tasks_result = await service.generate_photo_tasks(field_map)
    photo_tasks = tasks_result["photo_tasks"]

    # auto_judge: 接地電阻 = 120 Ω > 100 → fail
    judge_result = await service.auto_judge(
        field_name="接地電阻",
        measured_value=120,
        unit="Ω",
        equipment_type="低壓配電設備",
    )
    results.check(judge_result["judgment"] == "fail", "接地電阻 120Ω 判定為 fail",
                  f"judgment={judge_result['judgment']}")

    # 其他項目合格
    judge_pass = await service.auto_judge(
        field_name="絕緣電阻",
        measured_value=52.3,
        unit="MΩ",
        equipment_type="低壓配電設備",
    )
    results.check(judge_pass["judgment"] == "pass", "絕緣電阻 52.3MΩ 判定為 pass",
                  f"judgment={judge_pass['judgment']}")

    # 模擬 AI 結果 (接地電阻異常)
    anomaly_bindings = []
    for task in photo_tasks:
        display = task["display_name"]
        if "接地" in display:
            anomaly_bindings.append({
                "task_id": task["task_id"],
                "ai_result": {
                    "readings": [{"label": "接地電阻", "value": 120, "unit": "Ω"}],
                    "is_anomaly": True,
                    "condition_assessment": "不合格",
                    "anomaly_description": "接地電阻 120 Ohm 超過標準值 100 Ohm",
                    "summary": "接地電阻超標",
                },
                "value_field_ids": task["value_field_ids"],
                "judgment_field_ids": task["judgment_field_ids"],
                "remarks_field_ids": task["remarks_field_ids"],
            })
        else:
            anomaly_bindings.append({
                "task_id": task["task_id"],
                "ai_result": {
                    "readings": [{"label": display, "value": 50, "unit": "MΩ"}],
                    "is_anomaly": False,
                    "condition_assessment": "合格",
                    "anomaly_description": "",
                    "summary": "正常",
                },
                "value_field_ids": task["value_field_ids"],
                "judgment_field_ids": task["judgment_field_ids"],
                "remarks_field_ids": task["remarks_field_ids"],
            })

    results.check(len(anomaly_bindings) > 0, "已建立 anomaly bindings",
                  f"count={len(anomaly_bindings)}")

    # precision_map_fields
    mapping_result = await service.precision_map_fields(
        field_map=field_map,
        inspection_results=[{"equipment_name": "A棟配電盤"}],
        photo_task_bindings=anomaly_bindings,
    )
    mappings = mapping_result["mappings"]

    # 驗證不合格項有正確判定
    fail_mappings = [m for m in mappings
                     if '不合格' in str(m.get("suggested_value", ""))]
    results.check(len(fail_mappings) > 0, "有不合格判定",
                  f"fail_mappings={len(fail_mappings)}, all_values={[m.get('suggested_value','')[:30] for m in mappings]}")

    # 驗證備註有異常描述 (precision_map_fields writes anomaly_description into remarks)
    anomaly_mappings = [m for m in mappings
                        if '超過' in str(m.get("suggested_value", ""))
                        or '超標' in str(m.get("suggested_value", ""))
                        or 'Ohm' in str(m.get("suggested_value", ""))
                        or '120' in str(m.get("suggested_value", ""))]
    results.check(len(anomaly_mappings) > 0, "備註中有異常描述",
                  f"anomaly_mappings={len(anomaly_mappings)}, all_values={[m.get('suggested_value','')[:40] for m in mappings]}")

    # auto_fill 並驗證
    fill_values = [{"field_id": m["field_id"], "value": m["suggested_value"]}
                   for m in mappings]
    filled_bytes = await service.auto_fill(file_content, file_name, field_map, fill_values)
    results.check(len(filled_bytes) > 0, "auto_fill 產生非空結果")

    # batch_auto_judge 測試
    batch_results = await service.batch_auto_judge(
        readings=[
            {"field_name": "絕緣電阻 R相", "value": 52.3, "unit": "MΩ"},
            {"field_name": "接地電阻", "value": 120, "unit": "Ω"},
            {"field_name": "漏電斷路器動作時間", "value": 50, "unit": "ms"},
        ],
        equipment_type="低壓配電設備",
    )
    fail_count = sum(1 for r in batch_results if r["judgment"] == "fail")
    results.check(fail_count >= 1, "batch_auto_judge 正確偵測到 fail",
                  f"fail_count={fail_count}")

    return results


# ============================================================
# TEST 3: 消防設備定檢 — Word 格式
# ============================================================

async def test_case_3_fire_word():
    """消防設備定檢 — Word 格式"""
    print("\n" + "=" * 60)
    print("TEST 3: 消防設備定檢 — Word 格式")
    print("=" * 60)

    results = TestResults()
    service = FormFillService()

    file_content = create_fire_word_form()
    file_name = "消防設備定檢表.docx"

    # analyze_structure
    structure = await service.analyze_structure(file_content, file_name)
    results.check(structure["success"], "analyze_structure 成功")
    results.check(structure["file_type"] == "docx", "偵測為 docx 格式")
    results.check(structure["total_fields"] > 0, "偵測到 Word 欄位",
                  f"total_fields={structure['total_fields']}")

    field_map = structure["field_map"]

    # generate_photo_tasks
    tasks_result = await service.generate_photo_tasks(field_map)
    photo_tasks = tasks_result["photo_tasks"]
    results.check(len(photo_tasks) >= 0, "產生拍照任務或基本欄位",
                  f"tasks={len(photo_tasks)}, basic={len(tasks_result['basic_info_fields'])}")

    # auto_judge 消防項目
    judge_pressure = await service.auto_judge(
        field_name="滅火器壓力",
        measured_value=0.85,
        unit="MPa",
        equipment_type="滅火器",
    )
    results.check(judge_pressure["judgment"] == "pass", "滅火器壓力 0.85MPa 合格",
                  f"judgment={judge_pressure['judgment']}")

    # 模擬並回填
    fill_values = []
    for f in field_map:
        fid = f["field_id"]
        fname = f.get("field_name", "")
        if "滅火器" in fname:
            fill_values.append({"field_id": fid, "value": "0.85 MPa - 合格"})
        elif "緊急照明" in fname:
            fill_values.append({"field_id": fid, "value": "45 min - 合格"})
        elif "偵煙" in fname:
            fill_values.append({"field_id": fid, "value": "10%/m - 合格"})
        elif "消防栓" in fname:
            fill_values.append({"field_id": fid, "value": "2.0 kgf/cm2 - 合格"})
        elif "設備名稱" in fname or "名稱" in fname:
            fill_values.append({"field_id": fid, "value": "A棟消防設備"})
        elif "日期" in fname:
            fill_values.append({"field_id": fid, "value": "2026-03-13"})
        elif "人員" in fname:
            fill_values.append({"field_id": fid, "value": "李消防"})

    filled_bytes = await service.auto_fill(file_content, file_name, field_map, fill_values)
    results.check(len(filled_bytes) > 0, "Word auto_fill 產生非空結果")

    # 驗證 Word 可以正常載入
    try:
        doc = Document(io.BytesIO(filled_bytes))
        results.check(True, "回填後 Word 檔案可正常載入")

        # 檢查表格內容
        has_filled_content = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and '合格' in cell.text:
                        has_filled_content = True
                        break
        # Also check paragraphs
        for para in doc.paragraphs:
            if '2026' in para.text or '李消防' in para.text or 'A棟' in para.text:
                has_filled_content = True
                break
        results.check(has_filled_content, "Word 檔案含有回填內容")
    except Exception as e:
        results.check(False, "回填後 Word 載入", str(e))

    return results


# ============================================================
# TEST 4: 勾選式表單
# ============================================================

async def test_case_4_checkbox_form():
    """勾選式表單"""
    print("\n" + "=" * 60)
    print("TEST 4: 勾選式表單")
    print("=" * 60)

    results = TestResults()
    service = FormFillService()

    file_content = create_checkbox_excel_form()
    file_name = "勾選式檢查表.xlsx"

    # detect_checkbox_columns
    cb_result = await service.detect_checkbox_columns(file_content, file_name)
    dual_fields = cb_result.get("dual_column_fields", [])
    check_symbol = cb_result.get("check_symbol", "✓")

    results.check(len(dual_fields) > 0, "偵測到勾選雙欄結構",
                  f"fields={len(dual_fields)}")
    results.check(check_symbol is not None, "偵測到勾選符號",
                  f"symbol={check_symbol}")

    # 準備填入值 (混合合格/不合格)
    fill_values = []
    field_map = []
    for i, dual in enumerate(dual_fields):
        fid = dual["field_id"]
        if i == 2:  # 第三項設為不合格
            fill_values.append({
                "field_id": fid,
                "value": "不合格",
                "remarks": "需要重新測試",
            })
        else:
            fill_values.append({
                "field_id": fid,
                "value": "合格",
            })
        field_map.append({
            "field_id": fid,
            "field_name": dual["field_name"],
            "field_type": "dual_column_checkbox",
        })

    # auto_fill_with_checkboxes
    filled_bytes = await service.auto_fill_with_checkboxes(
        file_content=file_content,
        file_name=file_name,
        field_map=field_map,
        fill_values=fill_values,
        dual_column_fields=dual_fields,
        check_symbol=check_symbol,
    )
    results.check(len(filled_bytes) > 0, "auto_fill_with_checkboxes 產生結果")

    # 驗證回填內容
    wb = load_workbook(io.BytesIO(filled_bytes))
    ws = wb.active

    # 檢查有勾選符號被寫入
    has_check = False
    has_empty_opposite = False
    for row in range(2, ws.max_row + 1):
        b_val = ws.cell(row=row, column=2).value  # 合格欄
        c_val = ws.cell(row=row, column=3).value  # 不合格欄

        if b_val and str(b_val).strip():
            has_check = True
            # 對應不合格欄應為空
            if not c_val or not str(c_val).strip():
                has_empty_opposite = True
        if c_val and str(c_val).strip():
            has_check = True
            if not b_val or not str(b_val).strip():
                has_empty_opposite = True

    results.check(has_check, "勾選符號已寫入")
    results.check(has_empty_opposite, "互斥欄位正確(一欄有勾另一欄空)")

    # 驗證不合格項的備註
    has_remarks = False
    for row in range(2, ws.max_row + 1):
        d_val = ws.cell(row=row, column=4).value  # 備註欄
        if d_val and '重新測試' in str(d_val):
            has_remarks = True
            break
    results.check(has_remarks, "不合格項備註已填入")

    return results


# ============================================================
# TEST 5: 照片插入
# ============================================================

async def test_case_5_photo_insert():
    """照片插入"""
    print("\n" + "=" * 60)
    print("TEST 5: 照片插入")
    print("=" * 60)

    results = TestResults()
    service = FormFillService()

    # 建立測試照片
    photo1 = create_test_photo(200, 150, 'red')
    photo2 = create_test_photo(200, 150, 'blue')

    results.check(len(photo1) > 0, "測試照片 1 建立成功")
    results.check(len(photo2) > 0, "測試照片 2 建立成功")

    # ---- Excel 照片插入 ----
    excel_content = create_electrical_excel_form()
    photo_bindings = [
        {
            "task_id": "photo_001",
            "display_name": "絕緣電阻測量",
            "photo_bytes": photo1,
            "capture_time": "2026-03-13 14:30",
            "sequence": 1,
        },
        {
            "task_id": "photo_002",
            "display_name": "接地電阻測量",
            "photo_bytes": photo2,
            "capture_time": "2026-03-13 14:35",
            "sequence": 2,
        },
    ]

    excel_result = await service.insert_photos_into_report(
        file_content=excel_content,
        file_name="report.xlsx",
        photo_bindings=photo_bindings,
    )
    results.check(len(excel_result) > len(excel_content),
                  "Excel 插入照片後檔案變大",
                  f"before={len(excel_content)}, after={len(excel_result)}")

    # 驗證 Excel 可載入且有照片工作表
    wb = load_workbook(io.BytesIO(excel_result))
    photo_sheet_exists = any('照片' in name for name in wb.sheetnames)
    results.check(photo_sheet_exists, "Excel 有照片附件工作表",
                  f"sheets={wb.sheetnames}")

    # ---- Word 照片插入 ----
    word_content = create_fire_word_form()
    word_result = await service.insert_photos_into_report(
        file_content=word_content,
        file_name="report.docx",
        photo_bindings=photo_bindings,
    )
    results.check(len(word_result) > len(word_content),
                  "Word 插入照片後檔案變大",
                  f"before={len(word_content)}, after={len(word_result)}")

    # 驗證 Word 可正常載入
    try:
        doc = Document(io.BytesIO(word_result))
        results.check(True, "Word 含照片檔案可正常載入")
    except Exception as e:
        results.check(False, "Word 含照片載入", str(e))

    return results


# ============================================================
# TEST 6: 歷史資料帶入
# ============================================================

async def test_case_6_history_service():
    """歷史資料帶入"""
    print("\n" + "=" * 60)
    print("TEST 6: 歷史資料帶入")
    print("=" * 60)

    results = TestResults()

    # 使用臨時 DB
    tmp = tempfile.NamedTemporaryFile(suffix='.db', delete=False)
    tmp.close()
    history = HistoryService(db_path=tmp.name)

    equipment_id = "EQ-HIST-001"

    # 儲存第一次定檢
    h1 = await history.save_inspection(
        equipment_id=equipment_id,
        equipment_name="歷史測試設備",
        template_id="TEMP-001",
        inspection_date="2025-09-15",
        inspector="張三",
        results=[
            {"field_name": "絕緣電阻 R相", "value": 80, "unit": "MΩ", "judgment": "pass"},
            {"field_name": "接地電阻", "value": 55, "unit": "Ω", "judgment": "pass"},
        ],
    )
    results.check(h1 is not None, "第一次定檢儲存成功", f"history_id={h1}")

    # 儲存第二次定檢
    h2 = await history.save_inspection(
        equipment_id=equipment_id,
        equipment_name="歷史測試設備",
        template_id="TEMP-001",
        inspection_date="2025-12-15",
        inspector="張三",
        results=[
            {"field_name": "絕緣電阻 R相", "value": 65, "unit": "MΩ", "judgment": "pass"},
            {"field_name": "接地電阻", "value": 60, "unit": "Ω", "judgment": "pass"},
        ],
    )
    results.check(h2 is not None, "第二次定檢儲存成功")

    # 儲存第三次定檢
    h3 = await history.save_inspection(
        equipment_id=equipment_id,
        equipment_name="歷史測試設備",
        template_id="TEMP-001",
        inspection_date="2026-03-13",
        inspector="張三",
        results=[
            {"field_name": "絕緣電阻 R相", "value": 52, "unit": "MΩ", "judgment": "pass"},
            {"field_name": "接地電阻", "value": 65, "unit": "Ω", "judgment": "pass"},
        ],
    )
    results.check(h3 is not None, "第三次定檢儲存成功")

    # get_previous_values 驗證
    prev = await history.get_previous_values(
        equipment_id=equipment_id,
        field_names=["絕緣電阻 R相", "接地電阻"],
    )
    results.check("絕緣電阻 R相" in prev, "get_previous_values 回傳絕緣電阻",
                  f"keys={list(prev.keys())}")
    results.check(prev.get("絕緣電阻 R相", {}).get("value") == 52,
                  "前次絕緣電阻值正確 (52)",
                  f"value={prev.get('絕緣電阻 R相', {}).get('value')}")

    # 第二次定檢 — 驗證 get_history 回傳正確數量
    history_list = await history.get_history(equipment_id)
    results.check(len(history_list) == 3, "歷史記錄共 3 筆",
                  f"count={len(history_list)}")

    # 趨勢分析 — 絕緣電阻連續 3 次下降
    trend = await history.analyze_trend(
        equipment_id=equipment_id,
        field_name="絕緣電阻 R相",
        num_records=5,
    )
    results.check(trend is not None, "趨勢分析回傳結果")
    results.check(trend["trend"] == "declining", "趨勢為 declining",
                  f"trend={trend['trend']}")
    results.check(trend["consecutive_decline"] >= 2,
                  "連續下降次數 >= 2",
                  f"decline={trend['consecutive_decline']}")
    results.check(trend["warning"] is not None, "有警告訊息",
                  f"warning={trend.get('warning', '')[:50]}")
    results.check(len(trend["values"]) == 3, "趨勢值有 3 筆",
                  f"values={trend['values']}")

    # 清理
    try:
        os.unlink(tmp.name)
    except Exception:
        pass

    return results


# ============================================================
# 主程式
# ============================================================

async def main():
    print("=" * 60)
    print(" Sprint 6 Task 6.1: E2E 自動化測試")
    print("=" * 60)

    all_results = []

    test_functions = [
        ("Test 1: 電氣設備 — 全合格", test_case_1_electrical_all_pass),
        ("Test 2: 電氣設備 — 有異常", test_case_2_electrical_with_anomaly),
        ("Test 3: 消防設備 — Word", test_case_3_fire_word),
        ("Test 4: 勾選式表單", test_case_4_checkbox_form),
        ("Test 5: 照片插入", test_case_5_photo_insert),
        ("Test 6: 歷史資料帶入", test_case_6_history_service),
    ]

    for name, func in test_functions:
        try:
            result = await func()
            all_results.append((name, result))
        except Exception as e:
            print(f"\n  ERROR in {name}: {e}")
            import traceback
            traceback.print_exc()
            err_result = TestResults()
            err_result.check(False, f"{name} 執行失敗", str(e))
            all_results.append((name, err_result))

    # 彙總
    print("\n" + "=" * 60)
    print(" E2E 測試彙總")
    print("=" * 60)

    total_passed = 0
    total_failed = 0

    for name, result in all_results:
        status = "PASS" if result.failed == 0 else "FAIL"
        print(f"  [{status}] {name}: {result.passed} passed, {result.failed} failed")
        total_passed += result.passed
        total_failed += result.failed
        if result.errors:
            for err in result.errors:
                print(f"         -> {err}")

    print(f"\n  Total: {total_passed} passed, {total_failed} failed")
    rate = total_passed / (total_passed + total_failed) * 100 if (total_passed + total_failed) > 0 else 0
    print(f"  Pass Rate: {rate:.1f}%")

    return total_failed == 0


if __name__ == "__main__":
    success = asyncio.run(main())
    sys.exit(0 if success else 1)
